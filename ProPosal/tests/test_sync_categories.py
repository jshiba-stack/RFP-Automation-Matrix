"""Section I round-trip: updater._sync_categories rebuilds a clean, FY-standard
Categories table (sorted a-x, duplicates combined, canonical names, no red font).

Synthetic docx fixture + the committed FY2027 taxonomy — no reference FINAL needed.
"""

import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.table import _Cell

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))
from proposal import dit_taxonomy, updater  # noqa: E402
from proposal.flags import Report  # noqa: E402

TAX = dit_taxonomy.load_taxonomy()
BY = dit_taxonomy.by_letter(TAX)


def _cats_doc():
    """A Categories table whose one data row has a RED cell (annotation) to prove
    the rebuild clears color."""
    doc = Document()
    t = doc.add_table(rows=1, cols=3)
    for i, h in enumerate(("DIT #", "Professional Service Category", "Description")):
        t.rows[0].cells[i].text = h
    r = t.add_row()
    r.cells[0].text = "d"
    run = r.cells[1].paragraphs[0].add_run("Web Applications – Design and Development")
    run.font.color.rgb = RGBColor(0xEE, 0x00, 0x00)   # red working annotation
    r.cells[2].text = "JS"
    return doc, t


def _letters(t):
    return [t.rows[i].cells[0].text.strip() for i in range(1, len(t.rows))]


def _row_col0_vmerge(t, ri):
    """Read the vMerge state of a row's first <w:tc> directly (avoid python-docx
    returning the merge-origin cell for continuation rows)."""
    tc = t.rows[ri]._tr.tc_lst[0]
    tcPr = tc.find(qn("w:tcPr"))
    vm = tcPr.find(qn("w:vMerge")) if tcPr is not None else None
    if vm is None:
        return None
    return vm.get(qn("w:val")) or "continue"


def _row_col0_text(t, ri):
    return _Cell(t.rows[ri]._tr.tc_lst[0], t).text.strip()


def _has_explicit_color(t):
    for row in t.rows:
        for c in row.cells:
            for p in c.paragraphs:
                for run in p.runs:
                    rpr = run._r.find(qn("w:rPr"))
                    if rpr is not None and rpr.find(qn("w:color")) is not None:
                        return True
    return False


def test_sync_categories_finalizes_uppercase_xmerge_clears_color():
    doc, t = _cats_doc()
    store = [
        {"dit_number": "d", "name": BY["e"]["name"], "description": "JS"},   # stale d, name == e
        {"dit_number": "a", "name": BY["a"]["name"], "description": "Oracle"},
        {"dit_number": "x", "name": "Business Intelligence/AI", "description": "CoPilot"},
        {"dit_number": "?", "name": "IT Portfolio and Project Management", "description": "PMP"},
    ]
    rep = Report()
    updater._sync_categories(doc, store, rep)

    # the E row uses the canonical FY name (self-healed from stale d)
    e_row = next(i for i in range(1, len(t.rows)) if _row_col0_text(t, i) == "E")
    assert t.rows[e_row].cells[1].text.strip() == BY["e"]["name"]
    # X block: each specialty keeps its own title; column 1 is vertically merged
    x_rows = [i for i in range(1, len(t.rows)) if t.rows[i].cells[1].text.strip()
              in ("Business Intelligence/AI", "IT Portfolio and Project Management")]
    assert len(x_rows) == 2
    assert _row_col0_vmerge(t, x_rows[0]) == "restart"     # top of the span shows X
    assert _row_col0_vmerge(t, x_rows[1]) == "continue"    # merged into it
    assert _row_col0_text(t, x_rows[0]) == "X"
    assert _row_col0_text(t, x_rows[1]) == ""             # continuation cell is blank
    # the red annotation color is gone
    assert not _has_explicit_color(t)
    assert any("rebuilt Section I" in r.summary for r in rep.applied_records)


def test_sync_categories_preserves_description_line_breaks():
    doc, t = _cats_doc()
    store = [{"dit_number": "a", "name": BY["a"]["name"],
              "description": "Oracle\nMS SQL Server\nMongoDB"}]
    updater._sync_categories(doc, store, Report())
    a_row = next(i for i in range(1, len(t.rows)) if _row_col0_text(t, i) == "A")
    paras = [p.text for p in t.rows[a_row].cells[2].paragraphs]
    assert paras == ["Oracle", "MS SQL Server", "MongoDB"], "each item on its own line"


def test_sync_categories_missing_table_flags():
    doc = Document()
    doc.add_table(rows=1, cols=2)   # not a Categories signature
    rep = Report()
    updater._sync_categories(doc, [{"dit_number": "a", "name": "X"}], rep)
    assert any(f.kind == "MISSING" for f in rep.flags)
