"""Tests for the compliance checklist and the format checker."""

import sys
from pathlib import Path

import pytest
from docx import Document

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))
from proposal import compliance, formatcheck  # noqa: E402
from proposal.checks import FAIL, PASS, WARN  # noqa: E402

_REFS = ROOT / "assets/refs"
BASE = next(iter(sorted(_REFS.glob("*FINAL*.docx"))), _REFS / "__none__.docx")
BASE_PDF = next(iter(sorted(_REFS.glob("*FINAL*.pdf"))), _REFS / "__none__.pdf")
pytestmark = pytest.mark.skipif(not BASE.exists(), reason="reference FINAL not available")


def _status(rep, name_substr):
    for c in rep.checks:
        if name_substr.lower() in c.name.lower():
            return c.status
    return None


def test_checklist_all_pass_with_real_pdf():
    store = {"opportunity": {"selected_categories": ["1", "2"], "allowed_categories": ["1", "2", "3"]}}
    rep = compliance.run_checklist(BASE, store, {"pdf_size_cap_mb": 3.0, "page_limit": 30},
                                   pdf_path=BASE_PDF if BASE_PDF.exists() else None)
    assert _status(rep, "Required sections") == PASS
    assert _status(rep, "Categories") == PASS
    if BASE_PDF.exists():
        assert _status(rep, "PDF size") == PASS
        assert _status(rep, "Page limit") == PASS
        assert rep.ok


def test_checklist_warns_without_pdf():
    # pass an in-memory Document so no companion .pdf is auto-detected
    doc = Document(str(BASE))
    rep = compliance.run_checklist(doc, {}, {})
    assert _status(rep, "PDF size") == WARN
    assert _status(rep, "Page limit") == WARN


def test_checklist_fails_disallowed_category():
    store = {"opportunity": {"selected_categories": ["99"], "allowed_categories": ["1", "2"]}}
    rep = compliance.run_checklist(BASE, store, {})
    assert _status(rep, "Categories") == FAIL
    assert not rep.ok


def test_checklist_fails_missing_sections_on_blank_doc(tmp_path):
    p = tmp_path / "blank.docx"
    Document().save(str(p))
    rep = compliance.run_checklist(p, {}, {})
    assert _status(rep, "Required sections") == FAIL


def test_formatcheck_passes_on_final():
    rep = formatcheck.check_format(BASE, BASE)
    assert _status(rep, "Footer present") == PASS
    assert _status(rep, "Page numbering") == PASS
    assert _status(rep, "Heading styles") == PASS
    assert rep.ok


def test_formatcheck_fails_on_blank_doc(tmp_path):
    p = tmp_path / "blank.docx"
    Document().save(str(p))
    rep = formatcheck.check_format(p)
    assert _status(rep, "Footer present") == FAIL
    assert _status(rep, "Page numbering") == FAIL


if __name__ == "__main__":
    sys.exit(pytest.main([__file__, "-v"]))
