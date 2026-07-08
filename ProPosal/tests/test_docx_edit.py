"""Unit tests for docx_edit run-collapse replacement against REAL fragmentation.

Patterns lifted verbatim from the Phase-0 probe of the submittals:
  Fiscal Year -> ['Fiscal Year 202', '6']
  letter body -> [..., 'for the fiscal year ', '202', '6', '.  ']
  end-date     -> ['202', '5', '+']  /  ['202', '5+']
  footer       -> 'Acme - FY26 Professional Services Submittal' + tab + PAGE field
"""

import sys
from pathlib import Path

import pytest
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
from proposal import docx_edit  # noqa: E402


def _para_with_runs(run_texts, sizes=None):
    """Build a paragraph whose runs have the given texts (and optional sizes)."""
    doc = Document()
    p = doc.add_paragraph()
    for i, t in enumerate(run_texts):
        r = p.add_run(t)
        if sizes and sizes[i] is not None:
            r.font.size = sizes[i]
    return p


def test_fiscal_year_split_across_two_runs():
    p = _para_with_runs(["Fiscal Year 202", "6"])
    res = docx_edit.replace_in_paragraph(p, r"Fiscal Year 20\d\d", "Fiscal Year 2027")
    assert res.applied and res.matched
    assert docx_edit.para_text(p) == "Fiscal Year 2027"


def test_letter_body_year_only():
    runs = ["for the fiscal year ", "202", "6", ".  "]
    p = _para_with_runs(runs)
    # bump just the year digits, which live in two runs ('202' + '6')
    res = docx_edit.replace_in_paragraph(p, r"(?<=fiscal year )20\d\d", "2027")
    assert res.applied
    assert docx_edit.para_text(p) == "for the fiscal year 2027.  "


def test_end_date_ongoing_three_runs():
    p = _para_with_runs(["202", "5", "+"])
    res = docx_edit.replace_in_paragraph(p, r"20\d\d(?=\+)", "2026")
    assert res.applied
    assert docx_edit.para_text(p) == "2026+"


def test_end_date_ongoing_two_runs():
    p = _para_with_runs(["202", "5+"])
    res = docx_edit.replace_in_paragraph(p, r"20\d\d(?=\+)", "2026")
    assert res.applied
    assert docx_edit.para_text(p) == "2026+"


def test_completed_year_not_matched_by_ongoing_pattern():
    p = _para_with_runs(["202", "4"])  # '2024', no '+'
    res = docx_edit.replace_in_paragraph(p, r"20\d\d(?=\+)", "2026")
    assert not res.matched
    assert docx_edit.para_text(p) == "2024"


def test_only_touched_runs_change_field_preserved():
    """A footer-like paragraph: literal 'FY26' + a trailing PAGE field run.

    Replacing 'FY26' must leave the field run's element identity intact.
    """
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Acme - FY26 Professional Services Submittal\t")
    p.add_run("Page ")
    field_run = p.add_run("9")  # stand-in for the PAGE field result run
    field_el = field_run._element
    res = docx_edit.replace_in_paragraph(p, r"FY2\d", "FY27")
    assert res.applied
    assert res.runs_touched == 1
    assert docx_edit.para_text(p) == "Acme - FY27 Professional Services Submittal\tPage 9"
    # the field run object/element is the very same one (untouched)
    assert field_run._element is field_el
    assert field_run.text == "9"


def test_mixed_rpr_span_is_flagged_not_applied():
    """If the matched span crosses differing formatting, refuse + report."""
    from docx.shared import Pt

    p = _para_with_runs(["202", "6"], sizes=[Pt(12), Pt(24)])
    res = docx_edit.replace_in_paragraph(p, r"2026", "2027")
    assert res.matched and not res.applied
    assert res.rpr_uniform is False
    assert docx_edit.para_text(p) == "2026"  # unchanged


def test_replace_all_handles_repeats():
    p = _para_with_runs(["2026 and ", "202", "6 again"])
    out = docx_edit.replace_all_in_paragraph(p, r"2026", "2027")
    assert len(out) == 2
    assert docx_edit.para_text(p) == "2027 and 2027 again"


if __name__ == "__main__":
    sys.exit(pytest.main([__file__, "-v"]))
