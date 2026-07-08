"""Tests for the resume rebuild pipeline (resume_rebuild.py) — fictional data.

Covers: line extraction + respacing from a per-word-op PDF (the PDFescape
signature), structural parsing (sections, job/date pairing, bullets, wrapped
prose), the current-employer "Present" rule, docx rendering onto a template,
and the lost-words verification gate. Word-dependent export paths are not
run here (rebuild() end-to-end needs Word; its pieces are covered).
"""

import sys
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))
from proposal import resume_rebuild as RR  # noqa: E402
from test_resume_quality import HELVETICA, _raw_pdf  # noqa: E402


def _ln(y, x, text, bold=False, size=10.0, right=""):
    return RR.Line(y=y, x=x, size=size, bold=bold, text=text, right_text=right)


FICTION = [
    _ln(150, 72, "Jordan Avery", size=15, bold=True),
    _ln(165, 72, "Senior Widget Engineer", bold=True),
    _ln(190, 72, "SUMMARY", bold=True, size=11),
    _ln(210, 72, "Seasoned engineer with 10 years of experience in"),
    _ln(222, 72, "widget design and widget maintenance."),
    _ln(250, 72, "TECHNICAL SKILLS", bold=True),
    _ln(270, 72, "Languages: Widgetscript, COBOL"),
    _ln(282, 72, "Platforms: WidgetOS"),
    _ln(310, 72, "EXPERIENCE", bold=True),
    _ln(330, 72, "Acme Fictional, LLC, Springfield", bold=True,
        right="2020 to 2024"),
    _ln(344, 72, "Widget Lead"),
    _ln(370, 72, "Initech - Metropolis", bold=True, right="2015 to 2020"),
    _ln(384, 72, "Widget Engineer"),
    _ln(420, 72, "EDUCATION", bold=True),
    _ln(440, 90, "• State University - BS Widgetry"),
]


def test_parse_structure():
    p = RR.parse_resume(FICTION)
    assert p.name == "Jordan Avery"
    assert p.subtitle == "Senior Widget Engineer"
    headers = [h for h, _ in p.sections]
    assert headers == ["Summary", "Technical Skills", "Experience", "Education"]
    sec = dict(p.sections)
    # wrapped prose merged into ONE paragraph
    assert sec["Summary"] == [("para", "Seasoned engineer with 10 years of "
                               "experience in widget design and widget "
                               "maintenance.")]
    # "Label:" rows stay separate
    assert sec["Technical Skills"] == [("para", "Languages: Widgetscript, COBOL"),
                                       ("para", "Platforms: WidgetOS")]
    jobs = sec["Experience"]
    assert [(j.primary, j.dates, j.companion) for j in jobs] == [
        ("Acme Fictional, LLC, Springfield", "2020 to 2024", "Widget Lead"),
        ("Initech - Metropolis", "2015 to 2020", "Widget Engineer"),
    ]
    assert sec["Education"] == [("bullet", "State University - BS Widgetry")]


def test_footer_lines_dropped():
    lines = FICTION + [
        _ln(760, 72, "CONFIDENTIAL"),
        _ln(760, 300, "Jordan Avery - Widget Profile"),
        _ln(700, 72, "Page 1 of 1"),
    ]
    p = RR.parse_resume(lines)
    flat = str(p.sections)
    assert "CONFIDENTIAL" not in flat
    assert "Page 1 of 1" not in flat        # footer boilerplate, any position


def test_present_rule_rewrites_only_firm_entry():
    p = RR.parse_resume(FICTION)
    notes = RR.apply_present_rule(p, ["Acme Fictional, LLC"])
    jobs = dict(p.sections)["Experience"]
    assert jobs[0].dates == "2020 to Present"
    assert jobs[1].dates == "2015 to 2020"           # other employers untouched
    assert notes and "2020 to Present" in notes[0]
    # already-Present entries are left alone (no duplicate note)
    assert RR.apply_present_rule(p, ["Acme Fictional, LLC"]) == []


def test_build_docx_renders_house_structure(tmp_path):
    tpl = tmp_path / "tpl.docx"
    d = docx.Document()
    d.add_paragraph(style="Title").add_run("NAME HERE")
    d.add_paragraph("Job Title Here")
    d.save(str(tpl))

    p = RR.parse_resume(FICTION)
    RR.apply_present_rule(p, ["Acme Fictional, LLC"])
    out = RR.build_docx(p, tpl, tmp_path / "out.docx")

    got = docx.Document(str(out))
    texts = [(q.style.name, q.text) for q in got.paragraphs]
    assert texts[0] == ("Title", "JORDAN AVERY")               # uppercased
    assert texts[1][1] == "Senior Widget Engineer"
    assert ("Heading 1", "Summary") in texts
    assert ("Heading 2", "Acme Fictional, LLC, Springfield\t2020 to Present") in texts
    assert ("Heading 3", "Widget Lead") in texts
    assert ("List Paragraph", "State University - BS Widgetry") in texts
    # order: Experience heading before its first job
    idx = {t: i for i, t in enumerate(texts)}
    assert idx[("Heading 1", "Experience")] < idx[
        ("Heading 2", "Acme Fictional, LLC, Springfield\t2020 to Present")]


def _word_ops_content(words_xy, size=10):
    """One text op per word with NO space glyphs — the PDF-editor signature
    (a single BT block with a Tm reposition before every word, as PDFescape
    writes it)."""
    ops = [f"1 0 0 1 {x} {y} Tm ({w}) Tj" for w, x, y in words_xy]
    return (f"BT /F1 {size} Tf\n" + "\n".join(ops) + "\nET").encode()


def test_extract_lines_from_per_word_ops(tmp_path):
    pdf = _raw_pdf(tmp_path / "r.pdf", _word_ops_content([
        ("Jordan", 72, 640), ("Avery", 105, 640), ("builds", 133, 640),
        ("widgets", 164, 640),
    ]), {"/F1": HELVETICA})
    texts = [ln.text for ln in RR.extract_lines(pdf)]
    assert "Jordan Avery builds widgets" in texts


def _run(x, text, size=10.0, y=600.0, bold=False):
    return RR._Run(x=x, y=y, size=size, bold=bold, text=text)


def test_line_from_splits_date_column_only_on_real_gap():
    no_space = lambda t: None                      # force per-run fallback
    # prose reaching past x=380 must NOT be split (words are close together)
    prose = [_run(72 + i * 42, w) for i, w in enumerate(
        ["Decisive", "professional", "with", "years", "of", "experience",
         "in", "the", "areas", "of"])]
    ln = RR._line_from(prose, no_space)
    assert ln.right_text == ""
    assert ln.text.endswith("areas of")
    # a genuine date column across a wide gulf IS split
    job = [_run(72, "Acme"), _run(96, "Fictional,"), _run(140, "LLC"),
           _run(470, "2020"), _run(494, "to"), _run(506, "2024")]
    ln = RR._line_from(job, no_space)
    assert ln.text == "Acme Fictional, LLC"
    assert ln.right_text == "2020 to 2024"


def test_missing_words_gate(tmp_path):
    src = _raw_pdf(tmp_path / "src.pdf", _word_ops_content(
        [("alpha", 72, 600), ("bravo", 110, 600), ("charlie", 148, 600)]),
        {"/F1": HELVETICA})
    ok = _raw_pdf(tmp_path / "ok.pdf",
                  b"BT /F1 10 Tf 1 0 0 1 72 600 Tm (alpha bra vo charlie) Tj ET",
                  {"/F1": HELVETICA})   # split word: compact match must pass
    bad = _raw_pdf(tmp_path / "bad.pdf",
                   b"BT /F1 10 Tf 1 0 0 1 72 600 Tm (alpha charlie) Tj ET",
                   {"/F1": HELVETICA})
    assert RR.missing_words(src, ok) == set()
    assert RR.missing_words(src, bad) == {"bravo"}
    assert RR.missing_words(src, bad, ignore={"bravo"}) == set()


if __name__ == "__main__":
    import pytest
    sys.exit(pytest.main([__file__, "-v"]))
