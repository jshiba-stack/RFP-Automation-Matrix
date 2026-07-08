"""Unit tests for the table formatting-standard pass + output-name sanitizer.

Synthetic python-docx fixtures only (fictional data). The pass enforces the house
standard on every table:
  * font size auto-fixed to 12pt
  * text color auto-cleared to default black
  * border deviations from 0.5pt single flagged (not changed)
  * _safe_stem sanitizes user-supplied output names
"""

import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
from proposal import proofread  # noqa: E402
from proposal.flags import KIND_REVIEW, Report  # noqa: E402
from proposal.jobs import _safe_stem  # noqa: E402

TARGET = Pt(proofread.TABLE_FONT_PT)


def _silent(*_a, **_k):
    pass


def _sizes(table):
    return {run.font.size
            for row in table.rows
            for cell in row.cells
            for para in cell.paragraphs
            for run in para.runs
            if run.text.strip()}


def _table(doc, sizes):
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].paragraphs[0].add_run("Resource").font.size = TARGET
    t.rows[0].cells[1].paragraphs[0].add_run("Qualifications").font.size = TARGET
    for name, pt in sizes:
        row = t.add_row()
        row.cells[0].paragraphs[0].add_run(name).font.size = Pt(pt)
        row.cells[1].paragraphs[0].add_run("quals").font.size = Pt(pt)
    return t


def _house_borders(tbl):
    tblPr = tbl._tbl.tblPr
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in proofread._BORDER_EDGES:
        el = borders.makeelement(qn(f"w:{edge}"), {})
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), proofread._BORDER_SZ)
        borders.append(el)
    tblPr.append(borders)


def test_font_auto_normalized_to_house_size():
    doc = Document()
    t = _table(doc, [("Alex", 11), ("Blair", 9)])
    report = Report()
    out = proofread.proofread_document(doc, report, log=_silent)

    assert out["font_tables"] == 1
    assert _sizes(t) == {TARGET}, "every run should be forced to 12pt"
    assert any(r.status == "applied" and "12pt" in r.summary for r in report.applied_records)


def test_font_left_alone_when_already_house_size():
    doc = Document()
    t = _table(doc, [("Alex", proofread.TABLE_FONT_PT)])
    _house_borders(t)
    report = Report()
    out = proofread.proofread_document(doc, report, log=_silent)
    # pagination=1: the pass rightly marks the fresh table's rows no-split
    assert out == {"font_tables": 0, "color_tables": 0, "letterhead_lines": 0,
                   "pagination": 1, "pp_border_tables": 0, "border_flags": 0}
    # the only applied change is the pagination pass marking rows no-split
    assert all(r.location == "Pagination" for r in report.applied_records)
    assert not report.flags


def test_red_text_is_cleared_to_black():
    doc = Document()
    t = doc.add_table(rows=1, cols=1)
    run = t.rows[0].cells[0].paragraphs[0].add_run("? needs attention")
    run.font.size = TARGET
    run.font.color.rgb = RGBColor(0xEE, 0x00, 0x00)
    report = Report()
    out = proofread.proofread_document(doc, report, log=_silent)

    assert out["color_tables"] == 1
    rpr = run._r.find(qn("w:rPr"))
    assert rpr is None or rpr.find(qn("w:color")) is None, "explicit color removed"
    assert any(r.status == "applied" and "color" in r.summary for r in report.applied_records)


def test_nonstandard_borders_are_flagged_not_changed():
    doc = Document()
    t = _table(doc, [("Alex", proofread.TABLE_FONT_PT)])   # font ok, borders absent
    report = Report()
    out = proofread.proofread_document(doc, report, log=_silent)

    assert out["border_flags"] == 1
    assert any(r.is_flag and r.kind == KIND_REVIEW and "border" in r.summary
               for r in report.flags)
    # flag only — the table's border XML is untouched
    assert t._tbl.tblPr is None or t._tbl.tblPr.find(qn("w:tblBorders")) is None


def test_border_outliers_detector():
    doc = Document()
    good = _table(doc, [("Alex", proofread.TABLE_FONT_PT)])
    _house_borders(good)
    doc.add_table(rows=1, cols=1)   # plain table, no borders -> outlier "Table B"

    assert proofread.border_outliers(doc) == ["Table B"]


def _pastperf(doc, client):
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].paragraphs[0].add_run("Client")
    t.rows[0].cells[1].paragraphs[0].add_run(client)
    for label in ("Detailed Scope of Work", "Issue Resolution"):
        t.add_row().cells[0].paragraphs[0].add_run(label)
    return t


def test_section_iii_gets_full_cell_borders():
    doc = Document()
    t = _pastperf(doc, "Springfield Widget Foundation")
    report = Report()
    out = proofread.proofread_document(doc, report, log=_silent)

    assert out["pp_border_tables"] == 1
    # every cell now has all four borders (no gap between the last two rows)
    for row in t.rows:
        for cell in row.cells:
            tcb = cell._tc.find(qn("w:tcPr")).find(qn("w:tcBorders"))
            assert tcb is not None
            for edge in ("top", "bottom", "left", "right"):
                el = tcb.find(qn(f"w:{edge}"))
                assert el is not None and el.get(qn("w:val")) == "single"
    # a past-perf table is auto-bordered, never border-flagged
    assert not any("border" in r.summary for r in report.flags)


def test_safe_stem_sanitizes():
    assert _safe_stem("My Submittal.docx") == "My Submittal"
    assert _safe_stem('  "Acme FY26"  ') == "Acme FY26"
    assert _safe_stem("a/b:c*?.pdf") == "abc"          # ext stripped, illegal removed
    assert _safe_stem("") is None
    assert _safe_stem("   ") is None
    assert _safe_stem("***") is None                   # nothing usable left


# --- letterhead standard ------------------------------------------------------

def _header_doc(lines_with_fmt):
    """Doc whose section header holds letterhead-ish paragraphs.
    lines_with_fmt: [(text, color_or_None, size_pt_or_None)]"""
    doc = Document()
    hdr = doc.sections[0].header
    for text, color, size in lines_with_fmt:
        p = hdr.add_paragraph()
        r = p.add_run(text)
        if color is not None:
            r.font.color.rgb = color
        if size is not None:
            r.font.size = Pt(size)
    return doc


def test_letterhead_blue_firm_line_goes_black_10pt():
    doc = _header_doc([
        ("Acme Fictional, LLC", RGBColor(0x00, 0x70, 0xC0), 10),
        ("123 Example Rd, Ste 9", None, 9),
        ("Springfield, HI 96800", None, 9),
        ("www.example.test", None, 12),
    ])
    report = Report()
    fixed = proofread._normalize_letterhead(doc, report, _silent)
    assert fixed == 4
    runs = [p.runs[0] for p in doc.sections[0].header.paragraphs if p.runs]
    assert all(r.font.size == Pt(proofread.LETTERHEAD_FONT_PT) for r in runs)
    assert str(runs[0].font.color.rgb) == "000000"
    assert any(f"black {proofread.LETTERHEAD_FONT_PT:g}pt" in r.summary
               for r in report.applied_records)


def test_letterhead_ignores_non_letterhead_header_text():
    doc = _header_doc([("Draft for internal review", RGBColor(0xFF, 0, 0), 14)])
    report = Report()
    assert proofread._normalize_letterhead(doc, report, _silent) == 0
    r = doc.sections[0].header.paragraphs[-1].runs[0]
    assert r.font.size == Pt(14)          # untouched


def test_letterhead_idempotent():
    doc = _header_doc([("Acme Fictional, LLC", RGBColor(0, 0x70, 0xC0), 9)])
    report = Report()
    assert proofread._normalize_letterhead(doc, report, _silent) == 1
    assert proofread._normalize_letterhead(doc, report, _silent) == 0


def test_appendix_cover_centers_heading_and_section():
    from docx.oxml.ns import qn
    doc = Document()
    doc.add_paragraph("body text before")
    doc.add_paragraph("Appendix:  Resumes of Key Personnel", style="Heading 1")
    report = Report()
    changed = proofread._appendix_cover(doc, report, _silent)
    assert changed >= 2
    head = [p for p in doc.paragraphs if "Appendix" in p.text][0]
    assert head._p.get_or_add_pPr().find(qn("w:jc")).get(qn("w:val")) == "center"
    tail = doc.element.body.find(qn("w:sectPr"))
    assert tail.find(qn("w:vAlign")).get(qn("w:val")) == "center"
    prev = head._p.getprevious()
    assert prev.find(qn("w:pPr")).find(qn("w:sectPr")) is not None
    # earlier section stays top-aligned; pass is idempotent
    assert prev.find(qn("w:pPr")).find(qn("w:sectPr")).find(qn("w:vAlign")) is None
    assert proofread._appendix_cover(doc, report, _silent) == 0


def test_toc_entry_indent_hanging_half_inch():
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Inches
    doc = Document()
    try:
        doc.styles.add_style("toc 1", WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        pass
    report = Report()
    assert proofread._toc_entry_indent(doc, report, _silent) == 1
    pf = doc.styles["toc 1"].paragraph_format
    assert pf.left_indent == Inches(0.5)
    assert pf.first_line_indent == Inches(-0.5)
    # tab stops untouched: the dot-leader tab must be able to regenerate
    assert len(pf.tab_stops) == 0
    assert proofread._toc_entry_indent(doc, report, _silent) == 0  # idempotent
