"""Integration tests for the smart copy-and-update engine against the real FINAL.

The FY2026 FINAL + FY2027 deltas are the ground-truth acceptance target:
bumping it must change exactly the fiscal year (title + inline + footer), both
cover dates, and the ongoing Capacity end-dates -- nothing else.
"""

import sys
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))
from proposal import updater  # noqa: E402
from proposal.docx_edit import iter_all_paragraphs, para_text  # noqa: E402

BASE = next(iter(sorted((ROOT / "assets/refs").glob("*FINAL*.docx"))), ROOT / "assets/refs/__none__.docx")

pytestmark = pytest.mark.skipif(not BASE.exists(), reason="reference FINAL not available")


def _all_texts(doc):
    out = [para_text(p) for p in iter_all_paragraphs(doc)]
    for sec in doc.sections:
        for p in sec.footer.paragraphs:
            t = para_text(p)
            if t.strip():
                out.append("FOOTER: " + t)
    return out


def test_acceptance_exactly_ten_intended_changes():
    from docx import Document

    base_texts = _all_texts(Document(str(BASE)))
    doc, report = updater.build(BASE, {}, target_fy=2027, cover_date="2026-03-02")
    out_texts = _all_texts(doc)

    assert len(base_texts) == len(out_texts)
    changed = [(a, b) for a, b in zip(base_texts, out_texts) if a != b]
    assert len(changed) == 10, changed
    assert report.flags == []
    assert len(report.applied_records) == 10

    # the specific edits
    joined = "\n".join(out_texts)
    assert "Fiscal Year 2027" in joined
    assert "for the fiscal year 2027" in joined
    assert "March 2, 2026" in joined
    assert "FY27 Professional Services Submittal" in joined
    assert joined.count("2026+") == 5          # ongoing rows refreshed
    assert "2024" in joined and "2023" in joined  # completed years untouched
    assert "2025+" not in joined


def test_ongoing_uses_cover_date_year():
    doc, _ = updater.build(BASE, {}, target_fy=2028, cover_date="2027-01-15")
    texts = "\n".join(_all_texts(doc))
    assert "2027+" in texts and "2025+" not in texts


def test_fiscal_year_kept_without_store_or_override():
    """No more silent detected+1: the document keeps its own year by default."""
    from docx import Document

    base_joined = "\n".join(_all_texts(Document(str(BASE))))
    doc, report = updater.build(BASE, {}, cover_date="2026-03-02")
    joined = "\n".join(_all_texts(doc))
    detected = "Fiscal Year 2026" in base_joined
    assert detected and "Fiscal Year 2026" in joined and "Fiscal Year 2027" not in joined
    assert not [r for r in report.applied_records if "Fiscal year" in r.location]
    # the store's opportunity.fiscal_year still drives a bump when present
    doc2, _ = updater.build(BASE, {"opportunity": {"fiscal_year": 2027}}, cover_date="2026-03-02")
    assert "Fiscal Year 2027" in "\n".join(_all_texts(doc2))


def test_new_store_project_row_is_appended():
    store = {
        "projects": [
            {"id": "x", "client": "Brand New Client LLC", "project": "Net-New Thing",
             "start_year": 2026, "end": "ongoing"}
        ]
    }
    doc, report = updater.build(BASE, store, target_fy=2027, cover_date="2026-03-02")
    from proposal import docx_map
    tbl = docx_map.find_capacity_table(doc)
    last = tbl.rows[-1]
    assert last.cells[0].text.strip() == "Brand New Client LLC"
    assert last.cells[1].text.strip() == "Net-New Thing"
    assert last.cells[3].text.strip() == "2026+"        # ongoing renders as-of-year+
    assert any("added project row" in r.summary for r in report.applied_records)
    assert not [f for f in report.flags if f.kind == "ADD MANUALLY"]


def test_changed_store_project_updates_row_in_place():
    from docx import Document

    from proposal import docx_map
    base_tbl = docx_map.find_capacity_table(Document(str(BASE)))
    row1 = base_tbl.rows[1]
    client, project = row1.cells[0].text.strip(), row1.cells[1].text.strip()
    store = {"projects": [{"id": "x", "client": client, "project": project,
                           "start_year": 1999, "end": 2024}]}
    doc, report = updater.build(BASE, store, target_fy=2027, cover_date="2026-03-02")
    tbl = docx_map.find_capacity_table(doc)
    assert len(tbl.rows) == len(base_tbl.rows)          # updated, not appended
    assert tbl.rows[1].cells[2].text.strip() == "1999"
    assert tbl.rows[1].cells[3].text.strip() == "2024"
    assert sum("updated" in r.summary and "from store" in r.summary
               for r in report.applied_records) == 2


def test_new_pp_block_is_appended_with_cloned_formatting():
    from proposal import updater as _u
    store = {"past_performance": [{
        "id": "pp-new", "client": "Brand New Client LLC", "project": "Net-New Thing",
        "contact": "Pat Example", "phone": "(808) 555-0100",
        "scope": "Did the thing.", "issue_resolution": "No issues.",
    }]}
    doc, report = updater.build(BASE, store, target_fy=2027, cover_date="2026-03-02")
    blocks = _u._pp_blocks(doc)
    assert any(c == _u._norm_name("Brand New Client LLC") for _t, c, _p in blocks)
    new_tbl = blocks[-1][0]
    texts = {r.cells[0].text.strip().lower(): r.cells[1].text.strip() for r in new_tbl.rows}
    assert texts["project"] == "Net-New Thing"
    assert texts["detailed scope of work"] == "Did the thing."
    assert any("added engagement block" in r.summary for r in report.applied_records)


def test_projectless_pp_record_never_duplicates_a_block():
    """Old (client-only) store records must match an existing block for that
    client -- even when the client has several blocks -- not append a copy."""
    from docx import Document

    from proposal import updater as _u
    base_blocks = _u._pp_blocks(Document(str(BASE)))
    clients = [c for _t, c, _p in base_blocks]
    dup_client = next((c for c in clients if clients.count(c) > 1), clients[0])
    # reconstruct the display client from the doc for the store record
    tbl = next(t for t, c, _p in base_blocks if c == dup_client)
    display = tbl.rows[0].cells[1].text.replace("\n", " ").strip()
    store = {"past_performance": [{"id": "pp-x", "client": display}]}   # no project
    doc, report = updater.build(BASE, store, target_fy=2027, cover_date="2026-03-02")
    assert len(_u._pp_blocks(doc)) == len(base_blocks)      # nothing appended
    assert not any("added engagement block" in r.summary for r in report.applied_records)


def test_new_personnel_row_is_appended():
    from proposal import docx_map
    store = {"personnel": [{"id": "p-new", "name": "Brandnew Person",
                            "qualifications": "Certified example engineer."}]}
    doc, report = updater.build(BASE, store, target_fy=2027, cover_date="2026-03-02")
    tbl = docx_map.find_table_by_signature(doc, docx_map.SIG_QUALIFICATIONS)[0]
    assert tbl.rows[-1].cells[0].text.strip() == "Brandnew Person"
    assert tbl.rows[-1].cells[1].text.strip() == "Certified example engineer."
    assert any("added resource row" in r.summary for r in report.applied_records)


if __name__ == "__main__":
    sys.exit(pytest.main([__file__, "-v"]))
