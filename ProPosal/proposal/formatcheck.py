"""Final format checker (flag-only).

After a build or generate, confirm the formatting survived: the footer and its
page-number field are intact, and the section headings still use heading styles.
If a template is supplied, also confirm the output introduces no named styles the
template doesn't define (a sign a paste/edit dragged in foreign formatting).
"""

from __future__ import annotations

from docx import Document
from docx.document import Document as _DocClass

from . import docx_map, proofread
from .checks import ChecklistReport
from .docx_edit import para_text


def _as_doc(doc_or_path):
    return doc_or_path if isinstance(doc_or_path, _DocClass) else Document(str(doc_or_path))


def _has_page_field(doc) -> bool:
    for sec in doc.sections:
        if "PAGE" in sec.footer._element.xml or "PAGE" in sec.header._element.xml:
            return True
    return False


def _has_footer_text(doc) -> bool:
    for sec in doc.sections:
        if any(para_text(p).strip() for p in sec.footer.paragraphs):
            return True
    return False


def check_format(doc_or_path, template_or_path=None) -> ChecklistReport:
    doc = _as_doc(doc_or_path)
    rep = ChecklistReport("Format check")

    # footer + page number
    if _has_footer_text(doc):
        rep.pass_("Footer present", "running footer text found")
    else:
        rep.fail("Footer present", "no footer text in any section")
    if _has_page_field(doc):
        rep.pass_("Page numbering intact", "PAGE field present in footer/header")
    else:
        rep.fail("Page numbering intact", "no PAGE field found")

    # headings use a heading style. Prefer a heading/title-styled paragraph
    # (prose can mention a heading's text without being it); fall back to a
    # paragraph that *starts with* the heading, whose style is then reported.
    paras = [(para_text(p).strip(), p.style.name if p.style else "") for p in doc.paragraphs]
    bad_headings = []
    for h in docx_map.REQUIRED_HEADINGS:
        hl = h.lower()
        styled = next((s for t, s in paras
                       if hl in t.lower() and ("heading" in s.lower() or "title" in s.lower())), None)
        if styled is not None:
            continue
        match = next((s for t, s in paras if t.lower().startswith(hl)), None)
        if match is None:
            bad_headings.append(f"{h} (missing)")
        else:
            bad_headings.append(f"{h} (style={match!r})")
    if bad_headings:
        rep.fail("Heading styles intact", "; ".join(bad_headings))
    else:
        rep.pass_("Heading styles intact", "section headings use heading styles")

    # foreign styles vs template
    if template_or_path is not None:
        tdoc = _as_doc(template_or_path)
        tmpl_styles = {s.name for s in tdoc.styles}
        used = {p.style.name for p in doc.paragraphs if p.style}
        foreign = sorted(used - tmpl_styles)
        if foreign:
            rep.warn("No foreign styles", f"styles not in template: {', '.join(foreign)}")
        else:
            rep.pass_("No foreign styles", "all paragraph styles exist in the template")

    # table consistency (the proofread pass normalizes these on build; here we
    # just report the current state, so a standalone check still surfaces them)
    font_bad = proofread.font_outliers(doc)
    if font_bad:
        rep.warn("Table font size uniform", f"mixed sizes in: {', '.join(font_bad)}")
    else:
        rep.pass_("Table font size uniform", "each table's rows share one font size")
    border_bad = proofread.border_outliers(doc)
    if border_bad:
        rep.warn("Table borders consistent", f"missing interior border in: {', '.join(border_bad)}")
    else:
        rep.pass_("Table borders consistent", "sibling tables share interior borders")

    return rep
