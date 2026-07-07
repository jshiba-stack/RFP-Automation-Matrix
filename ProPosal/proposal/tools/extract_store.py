"""Bootstrap a data store from an existing submittal .docx.

Reads the Categories, Professional Qualifications, Capacity, and Past-Performance
tables out of a FINAL submittal and writes a YAML data store -- so you start from
your real content instead of hand-typing it, and generate mode has complete data.

Usage:
    python -m proposal.tools.extract_store "<FINAL.docx>" [-o data/stores/store.yaml]
"""

from __future__ import annotations

import argparse
import re
from collections import OrderedDict

import yaml
from docx import Document

from .. import docx_map
from ..docx_edit import para_text


def _slug(text: str, n: int = 24) -> str:
    s = re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-")
    return s[:n] or "item"


def _firm_from_cover(doc) -> dict:
    paras = [para_text(p).strip() for p in doc.paragraphs]
    firm = {"legal_name": "", "dba": "", "address_lines": [], "website": ""}
    for t in paras[:40]:
        if not t:
            continue
        if not firm["legal_name"] and re.search(r"\bInternational\b|\bLLC\b|Consultants", t) and "DBA" not in t:
            firm["legal_name"] = t
        elif t.upper().startswith("DBA") and not firm["dba"]:
            firm["dba"] = t[3:].strip()
        elif re.search(r"\bHI\b|Honolulu|Ste|PMB|Rd|Ave|Street|Suite", t) and "www" not in t.lower():
            if len(firm["address_lines"]) < 2 and re.search(r"\d", t):
                firm["address_lines"].append(t)
        elif "www." in t.lower() and not firm["website"]:
            firm["website"] = t
    return firm


def extract(path: str) -> OrderedDict:
    doc = Document(path)
    fy = docx_map.detect_fiscal_year(doc)
    store = OrderedDict()
    store["firm"] = _firm_from_cover(doc)
    store["opportunity"] = OrderedDict(
        agency="City and County of Honolulu",
        department="DIT",
        fiscal_year=fy,
        cover_date=None,
        required_form="general",
        page_limit=30,
        pdf_size_cap_mb=3.0,
        selected_categories=[],
    )

    # Categories (DIT # / Category / Description). Id is keyed off the category
    # NAME (stable + unique), not the DIT # -- a "?" or blank DIT # slugs to
    # nothing and would collide (every unlettered row -> "cat-item").
    cats = []
    used_ids: set[str] = set()
    for t in docx_map.find_table_by_signature(doc, docx_map.SIG_CATEGORIES):
        for row in t.rows[1:]:
            c = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if not any(c):
                continue
            num = c[0]
            # keep the description's line structure (each tech on its own line);
            # only the DIT # / name are flattened
            desc = row.cells[2].text.strip() if len(row.cells) > 2 else ""
            base = f"cat-{_slug(c[1] or num)}"
            cid, k = base, 2
            while cid in used_ids:
                cid, k = f"{base}-{k}", k + 1
            used_ids.add(cid)
            cats.append(OrderedDict(id=cid, dit_number=num,
                                    name=c[1], description=desc))
    store["categories"] = cats
    store["opportunity"]["selected_categories"] = [c["dit_number"] for c in cats if c["dit_number"]]

    # Personnel (Resource / Qualifications)
    people = []
    for t in docx_map.find_table_by_signature(doc, docx_map.SIG_QUALIFICATIONS):
        for row in t.rows[1:]:
            c = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if not c[0]:
                continue
            people.append(OrderedDict(id=_slug(c[0]), name=c[0],
                                      qualifications=c[1] if len(c) > 1 else "",
                                      resume_docx=f"resumes/{_slug(c[0])}.docx"))
    store["personnel"] = people

    # Projects (Capacity table)
    projects = []
    cap = docx_map.find_capacity_table(doc)
    if cap is not None:
        for row in cap.rows[1:]:
            c = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if not c[0]:
                continue
            end_raw = c[3] if len(c) > 3 else ""
            end = "ongoing" if end_raw.endswith("+") else (
                int(re.sub(r"\D", "", end_raw)) if re.search(r"\d", end_raw) else end_raw
            )
            start = c[2] if len(c) > 2 else ""
            start = int(re.sub(r"\D", "", start)) if re.search(r"\d", start) else start
            projects.append(OrderedDict(id=_slug(f"{c[0]}-{c[1]}"), client=c[0],
                                        project=c[1], start_year=start, end=end))
    store["projects"] = projects

    # Past performance: one record per 2-col block, all six fields (matched by
    # row label so the store can update/rebuild blocks later).
    label_to_field = {
        "client": "client", "project": "project", "client contact": "contact",
        "client phone": "phone", "detailed scope of work": "scope",
        "issue resolution": "issue_resolution",
    }
    pp = []
    for t in doc.tables:
        if not t.rows or len(t.rows[0].cells) != 2:
            continue
        if t.rows[0].cells[0].text.strip().lower() != docx_map.SIG_PASTPERF_FIRST_CELL:
            continue
        rec = OrderedDict()
        for row in t.rows:
            field = label_to_field.get(re.sub(r"\s+", " ", row.cells[0].text).strip().lower())
            if field:
                value = row.cells[1].text.strip()
                # client reads better on one line; scope keeps its line breaks
                rec[field] = value.replace("\n", " ") if field == "client" else value
        if rec.get("client"):
            rec_id = f"pp-{_slug(rec['client'])}"
            if rec.get("project"):
                rec_id += f"-{_slug(rec['project'], 16)}"
            pp.append(OrderedDict(id=rec_id, **rec))
    store["past_performance"] = pp

    return store


def _represent_ordereddict(dumper, data):
    return dumper.represent_mapping("tag:yaml.org,2002:map", data.items())


yaml.add_representer(OrderedDict, _represent_ordereddict)


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(prog="extract_store", description=__doc__)
    ap.add_argument("docx")
    ap.add_argument("-o", "--out", default="data/stores/store_extracted.yaml")
    args = ap.parse_args(argv)

    store = extract(args.docx)
    text = yaml.dump(store, sort_keys=False, allow_unicode=True, width=100)
    from pathlib import Path
    Path(args.out).parent.mkdir(parents=True, exist_ok=True)
    Path(args.out).write_text(text, encoding="utf-8")
    print(f"Wrote {args.out}")
    print(f"  categories={len(store['categories'])} personnel={len(store['personnel'])} "
          f"projects={len(store['projects'])} past_performance={len(store['past_performance'])}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
