"""Phase-0 read-only probe.

Dumps the structure of a submittal .docx so we can codify the anchor field map
in ``docx_map.py``: paragraph indices + styles, the *run* breakdown of the few
variable paragraphs (cover date / fiscal year), and every table's header-row
signature + dimensions. Run it against a new base whenever the template changes.

Usage:
    python -m proposal.tools.inspect_docx "<path-to.docx>" [--runs] [--tables]
"""

from __future__ import annotations

import re
import sys

from docx import Document
from docx.document import Document as _Doc
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph

# Paragraphs whose text trips one of these is "interesting" (likely variable).
INTEREST = re.compile(
    r"fiscal year|"
    r"[A-Z][a-z]+ \d{1,2}, \d{4}|"          # a Month D, YYYY date
    r"\b20\d{2}\b|"                          # any year
    r"dear|letter of interest",
    re.IGNORECASE,
)

HEADINGS = [
    "Categories",
    "Professional Qualifications",
    "Past Performance",
    "Capacity to Accomplish the Work",
    "Additional Criteria",
    "Appendix",
    "Table of Contents",
    "Professional Services Submittal Letter",
]


def _iter_block_items(parent):
    """Yield Paragraph and Table objects in document order."""
    body = parent.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def ptext(p: Paragraph) -> str:
    return "".join(r.text for r in p.runs)


def main() -> int:
    if len(sys.argv) < 2:
        print(__doc__)
        return 2
    path = sys.argv[1]
    show_runs = "--runs" in sys.argv

    doc: _Doc = Document(path)
    print(f"### {path}")

    # --- Block stream: paragraph index + tables in order ---
    print("\n## HEADINGS & INTERESTING PARAGRAPHS (paragraph index across body) ##")
    p_idx = 0
    t_idx = 0
    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            txt = ptext(block)
            is_head = any(h.lower() in txt.lower() for h in HEADINGS)
            if is_head or (txt.strip() and INTEREST.search(txt)):
                tag = "HEAD" if is_head else "var?"
                style = block.style.name if block.style else "?"
                print(f"  [p{p_idx:>3}] ({tag}) style={style!r:>22}  {txt[:80]!r}")
                if show_runs and (INTEREST.search(txt) or is_head):
                    runs = [r.text for r in block.runs]
                    print(f"          runs={runs}")
            p_idx += 1
        else:  # Table
            rows = block.rows
            ncols = len(block.columns)
            header = [c.text.strip() for c in rows[0].cells] if rows else []
            print(
                f"  [T{t_idx:>2}] table  {len(rows)}x{ncols}  "
                f"header={header}"
            )
            t_idx += 1

    # --- Capacity table detail: dump End-Date column run structure ---
    print("\n## TABLE CELL RUN DETAIL (cells containing a year) ##")
    for ti, tbl in enumerate(doc.tables):
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    txt = "".join(r.text for r in para.runs)
                    if re.search(r"20\d{2}", txt):
                        runs = [r.text for r in para.runs]
                        if len(runs) > 1 or "+" in txt or " " in txt.strip():
                            print(
                                f"  T{ti} r{ri} c{ci}: text={txt!r:>12}  runs={runs}"
                            )

    # --- Headers / footers ---
    print("\n## HEADERS / FOOTERS ##")
    for si, sec in enumerate(doc.sections):
        h = " | ".join(ptext(p) for p in sec.header.paragraphs if ptext(p).strip())
        f = " | ".join(ptext(p) for p in sec.footer.paragraphs if ptext(p).strip())
        has_pagefield = "PAGE" in sec.footer._element.xml or "PAGE" in sec.header._element.xml
        print(f"  section {si}: header={h!r}  footer={f!r}  page_field={has_pagefield}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
