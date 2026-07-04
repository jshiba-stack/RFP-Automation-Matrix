"""Structural anchors for a DIT/general-format submittal.

The historical .docx has no placeholder tokens, so we locate the variable
regions by *structure*: heading text, table header-row signatures, and the
position of date paragraphs relative to known headings. Constants here are the
field map confirmed by ``tools/inspect_docx.py`` against the FY2026 FINAL and
FY2027 DRAFT submittals.
"""

from __future__ import annotations

import re

from docx.document import Document as _Doc
from docx.table import Table
from docx.text.paragraph import Paragraph

from .docx_edit import iter_block_items, para_text

# Heading texts (matched case-insensitively, as a prefix/substring).
H_TOC = "Table of Contents"
H_LETTER = "Professional Services Submittal Letter"
H_CATEGORIES = "Categories"
H_QUALIFICATIONS = "Professional Qualifications"
H_PAST_PERF = "Past Performance"
H_CAPACITY = "Capacity to Accomplish the Work"
H_ADDITIONAL = "Additional Criteria"
H_APPENDIX = "Appendix"

REQUIRED_HEADINGS = [
    H_LETTER, H_CATEGORIES, H_QUALIFICATIONS, H_PAST_PERF,
    H_CAPACITY, H_ADDITIONAL, H_APPENDIX,
]

# Table header-row signatures (lowercased cell texts, first row).
SIG_CATEGORIES = ("dit #", "professional service category", "description")
SIG_QUALIFICATIONS = ("resource", "qualifications")
SIG_CAPACITY = ("client", "project", "start date", "end date")
SIG_PASTPERF_FIRST_CELL = "client"   # past-perf blocks are 2-col, first cell "Client"

# Regexes
RE_FULL_DATE = re.compile(r"[A-Z][a-z]+ \d{1,2}, \d{4}")
RE_FISCAL_YEAR = re.compile(r"Fiscal Year (20\d\d)")
RE_FISCAL_YEAR_INLINE = re.compile(r"(?<=fiscal year )(20\d\d)", re.IGNORECASE)
RE_FOOTER_FY = re.compile(r"FY(\d{2})")
RE_ONGOING_YEAR = re.compile(r"20\d\d(?=\s*\+)")   # year immediately before a '+'


def _is_heading(text: str, heading: str) -> bool:
    return heading.lower() in text.lower()


def block_index(doc: _Doc):
    """Return an ordered list of (kind, obj, p_idx, t_idx) for body blocks."""
    items = []
    p_idx = t_idx = 0
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            items.append(("p", block, p_idx, None))
            p_idx += 1
        else:
            items.append(("t", block, None, t_idx))
            t_idx += 1
    return items


def detect_fiscal_year(doc: _Doc) -> int | None:
    for p in doc.paragraphs:
        m = RE_FISCAL_YEAR.search(para_text(p))
        if m:
            return int(m.group(1))
    return None


def find_fiscal_year_paragraphs(doc: _Doc):
    """Paragraphs whose text contains a fiscal-year reference we should bump.

    Yields (paragraph, kind) where kind is 'fiscal_year' or 'inline'.
    """
    for p in doc.paragraphs:
        t = para_text(p)
        if RE_FISCAL_YEAR.search(t):
            yield p, "fiscal_year"
        elif RE_FISCAL_YEAR_INLINE.search(t):
            yield p, "inline"


def find_cover_date_paragraphs(doc: _Doc):
    """Every standalone date paragraph in the cover/letter region.

    The region is everything before the first 'Categories' heading (which opens
    the body). Within it we take only paragraphs whose *entire* trimmed text is a
    date -- the title-page date and the letter date -- so prose that merely
    mentions a date is never touched. Returns a list of (paragraph, label).
    """
    paras = [b[1] for b in block_index(doc) if b[0] == "p"]
    cat_i = next(
        (i for i, p in enumerate(paras) if _is_heading(para_text(p), H_CATEGORIES)),
        len(paras),
    )
    out = []
    n = 0
    for p in paras[:cat_i]:
        t = para_text(p).strip()
        if t and RE_FULL_DATE.fullmatch(t):
            n += 1
            label = "cover date (title page)" if n == 1 else "cover date (letter)"
            out.append((p, label))
    return out


def _table_signature(tbl: Table) -> tuple[str, ...]:
    if not tbl.rows:
        return ()
    return tuple(c.text.strip().lower() for c in tbl.rows[0].cells)


def find_table_by_signature(doc: _Doc, signature: tuple[str, ...]):
    """Find the first table whose header row equals ``signature`` (lowercased)."""
    matches = [t for t in doc.tables if _table_signature(t) == tuple(signature)]
    return matches


def find_capacity_table(doc: _Doc) -> Table | None:
    matches = find_table_by_signature(doc, SIG_CAPACITY)
    return matches[0] if matches else None


def find_past_performance_clients(doc: _Doc) -> dict[str, Table]:
    """Map client-name -> table for each 2-column 'Client' past-performance block."""
    out = {}
    for t in doc.tables:
        sig = _table_signature(t)
        if len(sig) == 2 and sig[0] == SIG_PASTPERF_FIRST_CELL:
            # header second cell holds the client name (may contain newlines)
            client = t.rows[0].cells[1].text.strip()
            out[client] = t
    return out


def footer_paragraphs(doc: _Doc):
    """Yield (paragraph, section_index) for every footer paragraph."""
    for si, sec in enumerate(doc.sections):
        for p in sec.footer.paragraphs:
            yield p, si


def _looks_like_heading_para(p: Paragraph, heading: str) -> bool:
    """True if ``p`` plausibly IS the heading (not prose that mentions it):
    either it uses a Heading/Title style, or its text starts with the heading."""
    t = para_text(p).strip()
    if not _is_heading(t, heading):
        return False
    style = (p.style.name if p.style else "").lower()
    return "heading" in style or "title" in style or t.lower().startswith(heading.lower())


def missing_headings(doc: _Doc) -> list[str]:
    missing = []
    for h in REQUIRED_HEADINGS:
        if not any(_looks_like_heading_para(p, h) for p in doc.paragraphs):
            missing.append(h)
    return missing
