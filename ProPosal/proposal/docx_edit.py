"""Low-level .docx editing helpers that survive Word's run fragmentation.

Word splits a single logical value across several ``w:r`` runs -- e.g.
``"Fiscal Year 2026"`` is stored as ``['Fiscal Year 202', '6']`` and an
end-date cell as ``['202', '5', '+']``. So a naive ``run.text.replace(...)``
never matches, because the target exists in no single run.

The helpers here match/replace at *paragraph (or table-cell) granularity*:
concatenate the run texts, find the span, then rewrite **only the runs the
match actually touches** -- leaving every other run untouched. That last point
matters: a footer paragraph contains a live ``PAGE`` field, and we must edit the
literal ``FY26`` text next to it without disturbing the field.

If the matched span crosses runs whose formatting (``rPr``) differs, collapsing
them into one run would silently change how the tail looks. We detect that and
refuse (``ReplaceResult.rpr_uniform == False``) so the caller can *flag* instead
of auto-applying.
"""

from __future__ import annotations

import copy
import re
from dataclasses import dataclass, field

from docx.document import Document as _Doc
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


# --- document traversal -----------------------------------------------------

def iter_block_items(parent):
    """Yield Paragraph and Table objects (top level) in document order."""
    body = parent.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def iter_all_paragraphs(doc: _Doc):
    """Yield every paragraph in the body, including inside table cells."""
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            yield block
        else:  # Table -- recurse one level (cells may hold nested tables, rare)
            for row in block.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        yield para


def para_text(para: Paragraph) -> str:
    """Concatenated visible text of a paragraph's runs."""
    return "".join(r.text for r in para.runs)


# --- replacement ------------------------------------------------------------

@dataclass
class ReplaceResult:
    matched: bool = False
    old: str = ""
    new: str = ""
    rpr_uniform: bool = True       # False -> span crossed differing formatting
    applied: bool = False
    runs_touched: int = 0


def _rpr_signature(run) -> str:
    """A cheap, stable signature of a run's formatting for uniformity checks."""
    rpr = run._element.find(qn("w:rPr"))
    return "" if rpr is None else str(rpr.xml)


def replace_in_paragraph(
    para: Paragraph,
    pattern: str | re.Pattern,
    repl: str,
    *,
    flags: int = 0,
    apply_when_mixed: bool = False,
) -> ReplaceResult:
    """Replace the first regex match of ``pattern`` in ``para``'s text.

    ``repl`` may reference groups (``\\1`` / ``\\g<name>``) -- it is expanded
    against the match. Only the runs the match touches are modified. If the span
    crosses runs with differing ``rPr`` the replacement is skipped unless
    ``apply_when_mixed`` is True; either way ``rpr_uniform`` reports the truth.
    """
    rx = re.compile(pattern, flags) if isinstance(pattern, str) else pattern
    runs = list(para.runs)
    if not runs:
        return ReplaceResult()
    texts = [r.text for r in runs]
    full = "".join(texts)
    m = rx.search(full)
    if not m:
        return ReplaceResult()

    start, end = m.start(), m.end()
    new_text = m.expand(repl)

    # Map character offsets -> (run index, offset within run).
    bounds = []  # (run_index, run_start, run_end)
    pos = 0
    for i, t in enumerate(texts):
        bounds.append((i, pos, pos + len(t)))
        pos += len(t)

    first_i = next(i for i, s, e in bounds if s <= start < e or (s == e == start))
    # last touched run = run containing end-1 (end is exclusive)
    last_i = next(i for i, s, e in bounds if s <= end - 1 < e)

    # uniformity check across touched runs (empty runs render nothing, so
    # their rPr can't change how the text looks -- ignore them)
    sigs = {_rpr_signature(runs[i]) for i in range(first_i, last_i + 1) if runs[i].text}
    uniform = len(sigs) <= 1

    result = ReplaceResult(
        matched=True,
        old=m.group(0),
        new=new_text,
        rpr_uniform=uniform,
        runs_touched=last_i - first_i + 1,
    )
    if not uniform and not apply_when_mixed:
        return result  # matched but not applied -> caller flags

    f_s = bounds[first_i][1]
    l_s = bounds[last_i][1]
    head = texts[first_i][: start - f_s]
    tail = texts[last_i][end - l_s :]

    if first_i == last_i:
        runs[first_i].text = head + new_text + tail
    else:
        runs[first_i].text = head + new_text
        for i in range(first_i + 1, last_i):
            runs[i].text = ""
        runs[last_i].text = tail

    result.applied = True
    return result


def replace_all_in_paragraph(
    para: Paragraph, pattern, repl, *, flags: int = 0, apply_when_mixed: bool = False
) -> list[ReplaceResult]:
    """Repeatedly replace until no match remains (guards against loops)."""
    out = []
    for _ in range(50):
        res = replace_in_paragraph(
            para, pattern, repl, flags=flags, apply_when_mixed=apply_when_mixed
        )
        if not res.matched:
            break
        out.append(res)
        if not res.applied:
            break  # mixed-rPr and not applying -> stop to avoid infinite loop
    return out


# --- table construction (generate mode) ------------------------------------

def set_cell_text(cell: _Cell, text: str) -> None:
    """Replace a cell's text while keeping its first paragraph/run formatting.

    The first paragraph's first run takes the new text; any extra runs are
    cleared and any extra paragraphs in the cell are removed.
    """
    text = "" if text is None else str(text)
    para = cell.paragraphs[0]
    # drop extra paragraphs (keep the first, which carries the style)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    if para.runs:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(text)


def append_cloned_row(table: Table, values: list[str], *, model_row: int = -1):
    """Append one data row, cloning an existing row's formatting/borders.

    ``model_row`` indexes the existing data rows (default: the last one).
    Returns the new row. Values are padded/truncated to the column count.
    """
    data_rows = table.rows[1:]
    if not data_rows:
        raise ValueError("table has no data row to use as a formatting model")
    new_tr = copy.deepcopy(data_rows[model_row]._tr)
    table._tbl.append(new_tr)
    new_row = table.rows[-1]
    ncols = len(table.columns)
    padded = list(values)[:ncols] + [""] * (ncols - len(values))
    for ci, val in enumerate(padded):
        set_cell_text(new_row.cells[ci], val)
    return new_row


def rebuild_table_body(table: Table, rows: list[list[str]], *, model_row: int = 1) -> int:
    """Replace a table's data rows with ``rows``, cloning a model row's format.

    Row 0 (header) is preserved. ``model_row`` is a 1-based index into the data
    rows whose cell formatting/borders/shading is cloned for every new row.
    Returns the number of rows written. Each row in ``rows`` is padded/truncated
    to the table's column count.
    """
    data_rows = table.rows[1:]
    if not data_rows:
        raise ValueError("table has no data row to use as a formatting model")
    ncols = len(table.columns)
    idx = min(max(model_row, 1), len(data_rows)) - 1
    model_xml = copy.deepcopy(data_rows[idx]._tr)

    for r in data_rows:                       # remove existing data rows
        r._tr.getparent().remove(r._tr)

    for values in rows:
        new_tr = copy.deepcopy(model_xml)
        table._tbl.append(new_tr)
        new_row = table.rows[-1]
        padded = list(values)[:ncols] + [""] * (ncols - len(values))
        for ci, val in enumerate(padded):
            set_cell_text(new_row.cells[ci], val)
    return len(rows)
