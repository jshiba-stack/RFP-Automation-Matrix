"""Formatting-standard pass (the submittal house style): tables + letterhead.

Runs after the build/generate engines and before save. Enforces the firm's
table standard across **every table** in the document:

* **font size -> 12pt** (auto-fixed): imported/edited content drifts a point or
  two; every cell is brought to the house size;
* **text color -> default black** (auto-fixed): working drafts get red/coloured
  annotations that must not ship;
* **borders -> 0.5pt single on all edges** (flagged, not changed): border work is
  structural and easy to get wrong automatically, so deviations are raised as a
  REVIEW flag for a human rather than rewritten.

Page-header **letterhead** lines (firm name/address/site) are also normalized:
black, ``LETTERHEAD_FONT_PT``, right-aligned with their textbox pinned to the
right margin -- see ``_normalize_letterhead``.

"Auto-fix the safe stuff, flag the structural stuff." Section I additionally has
its own structural standard (uppercase A-X, canonical names, merged catch-all),
applied earlier by ``updater.write_finalized_categories``.

The module also exposes read-only detectors (``font_outliers``,
``border_outliers``) used by the standalone format checker to report state.
"""

from __future__ import annotations

import re

from docx.oxml.ns import qn
from docx.shared import Pt

from . import docx_map
from .flags import KIND_REVIEW

# --- house standard ---------------------------------------------------------

TABLE_FONT_PT = 12.0     # every table cell renders at this size
BORDER_PT = 0.5          # every table border is a single line this wide (half-point)
LETTERHEAD_FONT_PT = 9.0   # firm letterhead in page headers: black, this size
# (9pt: contact blocks conventionally sit below body size so they read as
# quiet letterhead; the firm's 2026-era resume exports already used 9pt)

_TARGET_SIZE = Pt(TABLE_FONT_PT)
_TARGET_EMU = int(_TARGET_SIZE)
_BORDER_SZ = str(int(round(BORDER_PT * 8)))   # Word stores width in 1/8 pt: 0.5pt -> "4"
_BORDER_EDGES = ("top", "left", "bottom", "right", "insideH", "insideV")

#: A page-header line that reads as firm letterhead (name w/ legal suffix,
#: street address, city+zip, url/email). Deliberately generic -- no firm data.
_LETTERHEADISH = re.compile(
    r"www\.|https?://|@"
    r"|\b\d{5}(-\d{4})?\b"
    r"|\b(LLC|L\.L\.C\.|Inc\.?|Ltd\.?|Corp\.?)(\b|$)"
    r"|\b(P\.?O\.?\s?Box|Suite|Ste|PMB|Rd|Road|St|Street|Ave|Avenue|Blvd)\b",
    re.I)


# --- labels -----------------------------------------------------------------

def _table_label(tbl, ordinal: int) -> str:
    """Human-readable place for a table, keyed off its header signature."""
    sig = docx_map._table_signature(tbl)
    letter = chr(ord("A") + ordinal) if ordinal < 26 else f"#{ordinal + 1}"
    if sig == docx_map.SIG_CATEGORIES:
        return "Section I · Categories"
    if sig == docx_map.SIG_QUALIFICATIONS:
        return "Section II · Qualifications"
    if sig == docx_map.SIG_CAPACITY:
        return "Section IV · Capacity"
    if len(sig) == 2 and sig[0] == docx_map.SIG_PASTPERF_FIRST_CELL:
        client = tbl.rows[0].cells[1].text.strip().splitlines()[0] if tbl.rows else ""
        return f"Section III · Past performance ({client})" if client \
            else f"Section III · Past performance (Table {letter})"
    return f"Table {letter}"


# --- font size --------------------------------------------------------------

def _normal_size(doc):
    try:
        return doc.styles["Normal"].font.size
    except (KeyError, AttributeError):
        return None


def _effective_size(run, para, normal_size):
    """The size a run actually renders at: explicit run size, else the nearest
    paragraph-style size up the base-style chain, else the Normal default."""
    if run.font.size is not None:
        return run.font.size
    style = para.style
    seen = 0
    while style is not None and seen < 10:
        if getattr(style, "font", None) is not None and style.font.size is not None:
            return style.font.size
        style = getattr(style, "base_style", None)
        seen += 1
    return normal_size


def _at_target_size(eff) -> bool:
    return eff is not None and int(eff) == _TARGET_EMU


def _normalize_font_size(doc, report, log) -> int:
    """Force every table cell to the house font size (auto-fix)."""
    normal = _normal_size(doc)
    changed = 0
    for ordinal, tbl in enumerate(doc.tables):
        fixed = 0
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if not run.text.strip() or _at_target_size(_effective_size(run, para, normal)):
                            continue
                        run.font.size = _TARGET_SIZE
                        fixed += 1
        if fixed:
            changed += 1
            report.applied(_table_label(tbl, ordinal),
                           f"normalized {fixed} run(s) to {TABLE_FONT_PT:g}pt")
            log(f"[proofread] {_table_label(tbl, ordinal)}: {fixed} run(s) -> {TABLE_FONT_PT:g}pt")
    return changed


# --- text color -------------------------------------------------------------

def _clear_colors(doc, report, log) -> int:
    """Remove explicit run colors in every table so text uses the default black."""
    changed = 0
    for ordinal, tbl in enumerate(doc.tables):
        cleared = 0
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        rpr = run._r.find(qn("w:rPr"))
                        if rpr is None:
                            continue
                        col = rpr.find(qn("w:color"))
                        if col is not None:
                            rpr.remove(col)
                            cleared += 1
        if cleared:
            changed += 1
            report.applied(_table_label(tbl, ordinal),
                           f"cleared {cleared} colored run(s) to default black")
            log(f"[proofread] {_table_label(tbl, ordinal)}: cleared {cleared} color(s)")
    return changed


# --- borders ----------------------------------------------------------------
# Section III standard: every past-performance cell is fully bordered (the firm
# does this by hand -- click the cell, "all borders"). We do it for them, so a
# built block never ships with a gap between rows. tcBorders must precede these
# tcPr children to stay schema-valid.
_TCBORDERS_AFTER = [qn(f"w:{t}") for t in (
    "shd", "noWrap", "tcMar", "textDirection", "tcFitText", "vAlign", "hideMark")]


def _is_pastperf(tbl) -> bool:
    sig = docx_map._table_signature(tbl)
    return len(sig) == 2 and sig[0] == docx_map.SIG_PASTPERF_FIRST_CELL


def _set_full_cell_borders(cell) -> None:
    """Give a cell a single 0.5pt border on all four edges (replacing any existing
    tcBorders), so the whole table renders a complete grid."""
    tcPr = cell._tc.get_or_add_tcPr()
    old = tcPr.find(qn("w:tcBorders"))
    if old is not None:
        tcPr.remove(old)
    tcb = tcPr.makeelement(qn("w:tcBorders"), {})
    for edge in ("top", "left", "bottom", "right"):
        el = tcb.makeelement(qn(f"w:{edge}"), {})
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), _BORDER_SZ)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tcb.append(el)
    anchor = next((c for c in tcPr if c.tag in _TCBORDERS_AFTER), None)
    (anchor.addprevious if anchor is not None else tcPr.append)(tcb)


def _apply_pastperf_borders(doc, report, log) -> int:
    """Auto-fix: full cell borders on every Section III past-performance table."""
    fixed = 0
    for ordinal, tbl in enumerate(doc.tables):
        if not _is_pastperf(tbl):
            continue
        for row in tbl.rows:
            for cell in row.cells:
                _set_full_cell_borders(cell)
        fixed += 1
        report.applied(_table_label(tbl, ordinal),
                       "added full cell borders (Section III standard)")
        log(f"[proofread] {_table_label(tbl, ordinal)}: full cell borders")
    return fixed


def _table_borders_ok(tbl) -> bool:
    """True if the table has a single 0.5pt border on all six edges."""
    tblPr = tbl._tbl.tblPr
    borders = tblPr.find(qn("w:tblBorders")) if tblPr is not None else None
    if borders is None:
        return False
    for edge in _BORDER_EDGES:
        el = borders.find(qn(f"w:{edge}"))
        if el is None or el.get(qn("w:val")) != "single" or el.get(qn("w:sz")) != _BORDER_SZ:
            return False
    return True


def _flag_borders(doc, report, log) -> int:
    """Flag every table whose borders aren't the house 0.5pt single (no auto-fix).

    Section III past-performance tables are excluded -- they are auto-bordered by
    ``_apply_pastperf_borders`` instead.
    """
    flagged = 0
    for ordinal, tbl in enumerate(doc.tables):
        if _is_pastperf(tbl) or _table_borders_ok(tbl):
            continue
        report.flag(_table_label(tbl, ordinal),
                    f"borders are not {BORDER_PT:g}pt single on all edges — set by hand",
                    KIND_REVIEW)
        flagged += 1
    return flagged


# --- read-only detectors (for the format checker) ---------------------------

def _table_font_ok(tbl, normal) -> bool:
    return not any(
        run.text.strip() and not _at_target_size(_effective_size(run, para, normal))
        for row in tbl.rows for cell in row.cells
        for para in cell.paragraphs for run in para.runs)


def font_outliers(doc) -> list[str]:
    """Labels of tables with any cell not at the house font size."""
    normal = _normal_size(doc)
    return [_table_label(t, o) for o, t in enumerate(doc.tables)
            if not _table_font_ok(t, normal)]


def border_outliers(doc) -> list[str]:
    """Labels of non-past-performance tables whose borders aren't the house 0.5pt."""
    return [_table_label(t, o) for o, t in enumerate(doc.tables)
            if not _is_pastperf(t) and not _table_borders_ok(t)]


# --- entry point ------------------------------------------------------------

def _normalize_letterhead(doc, report, log) -> int:
    """Letterhead standard: every letterhead-looking line in a page header
    renders black at LETTERHEAD_FONT_PT -- no blue firm names, no size drift.

    Walks every header part's paragraphs at the XML level so lines inside
    textboxes (where letterheads usually live, incl. both mc:AlternateContent
    copies) are covered. Auto-fixed: color and size are cosmetic here.

    Alignment is part of the standard too: letterhead lines are right-aligned
    and their textbox (if any) is anchored to the right MARGIN with no right
    inset, so the block's longest line ends flush with the body text's right
    edge instead of drifting past or short of it.
    """
    fixed = 0
    seen: set[int] = set()
    for sec in doc.sections:
        hdr = sec.header
        if id(hdr._element) in seen:      # linked headers share the part
            continue
        seen.add(id(hdr._element))
        for p_el in hdr._element.iter(qn("w:p")):
            # DIRECT runs only: a paragraph that merely CONTAINS a letterhead
            # textbox (plus, typically, the inline logo image) must not be
            # touched -- right-aligning it would move the logo.
            text = "".join(t.text or ""
                           for r in p_el.findall(qn("w:r"))
                           for t in r.findall(qn("w:t"))).strip()
            if not text or not _LETTERHEADISH.search(text):
                continue
            changed = _right_align_letterhead(p_el)
            for r_el in p_el.iter(qn("w:r")):
                if r_el.find(qn("w:t")) is None:
                    continue
                rpr = r_el.find(qn("w:rPr"))
                if rpr is None:
                    rpr = r_el.makeelement(qn("w:rPr"), {})
                    r_el.insert(0, rpr)
                color = rpr.find(qn("w:color"))
                if color is not None and color.get(qn("w:val")) not in (
                        "000000", "auto"):
                    color.set(qn("w:val"), "000000")
                    changed = True
                sz_val = str(int(LETTERHEAD_FONT_PT * 2))
                for tag in ("w:sz", "w:szCs"):
                    el = rpr.find(qn(tag))
                    if el is None:
                        el = rpr.makeelement(qn(tag), {})
                        rpr.append(el)
                        el.set(qn("w:val"), sz_val)
                        changed = True
                    elif el.get(qn("w:val")) != sz_val:
                        el.set(qn("w:val"), sz_val)
                        changed = True
            if changed:
                fixed += 1
                report.applied("Letterhead",
                               f"header line -> black {LETTERHEAD_FONT_PT:g}pt, "
                               f"right-aligned to margin: \"{text[:40]}\"")
    return fixed


_WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
_WPS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"

# --- pagination standard ------------------------------------------------------
# Cleaner page flow for the deliverable:
#   * table rows never split across pages (a row that doesn't fit moves whole);
#   * the Section IV heading and the Appendix divider always start a new page;
#   * Section III past-performance tables use a narrower label column so the
#     description column wraps less (shorter tables, fewer page spills).
_BREAK_BEFORE_HEADINGS = ("capacity to accomplish",)
PASTPERF_LABEL_COL_TWIPS = 2160     # 1.50" (was 1.68"); total table width kept

#: sectPr children that must FOLLOW vAlign (OOXML schema order)
_SECTPR_AFTER_VALIGN = ("noEndnote", "titlePg", "textDirection", "bidi",
                        "rtlGutter", "docGrid", "printerSettings")


def _toc_entry_indent(doc, report, log) -> int:
    """Uniform TOC entries: Word tabs each entry's text to the next default
    0.25\" stop after the heading numeral, so wide numerals (III., IV., VI.)
    land one stop further than narrow ones. A hanging indent on the TOC 1
    style pins every entry's text at 0.5\" (the numeral lives in the hang)
    WITHOUT touching tab stops -- the entries' dot-leader page-number tab is
    regenerated by Word at update time and must survive."""
    from docx.shared import Inches

    style = None
    for name in ("toc 1", "TOC 1"):
        try:
            style = doc.styles[name]
            break
        except KeyError:
            continue
    if style is None:
        return 0
    pf = style.paragraph_format
    if pf.left_indent == Inches(0.5) and pf.first_line_indent == Inches(-0.5):
        return 0
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.5)
    report.applied("Pagination",
                   "TOC entries aligned (0.5\" hanging indent after the numeral)")
    return 1


def _appendix_cover(doc, report, log) -> int:
    """The Appendix divider renders as a cover page for the resume pages:
    heading centered horizontally AND vertically. The heading is the last
    content in the document, so it gets its own final section (a section
    break is inserted before it) with vertical alignment = center."""
    from copy import deepcopy

    target = None
    for para in doc.paragraphs:
        if (para.style.name.startswith("Heading 1")
                and "appendix" in para.text.strip().lower()[:12]):
            target = para
            break
    if target is None:
        return 0
    changed = 0

    ppr = target._p.get_or_add_pPr()
    jc = ppr.find(qn("w:jc"))
    if jc is None:
        mark_rpr = ppr.find(qn("w:rPr"))
        jc = ppr.makeelement(qn("w:jc"), {})
        (mark_rpr.addprevious if mark_rpr is not None else ppr.append)(jc)
    if jc.get(qn("w:val")) != "center":
        jc.set(qn("w:val"), "center")
        changed += 1


    body = doc.element.body
    tail_sectpr = body.find(qn("w:sectPr"))
    if tail_sectpr is None:
        return changed

    # the manual page break that used to give the divider its page is
    # superseded by the section break and would leave a blank page
    el = target._p.getprevious()
    while el is not None and el.tag == qn("w:p") and not "".join(
            t.text or "" for t in el.iter(qn("w:t"))).strip():
        for br in list(el.iter(qn("w:br"))):
            if br.get(qn("w:type")) == "page":
                br.getparent().remove(br)
                changed += 1
        el = el.getprevious()

    # normalize the empty paragraphs after the heading to EXACTLY four: they
    # are part of the vertically-centered block, so four ~14pt lines below
    # the heading lift it ~2 lines above true center (the base doc's dozen
    # stray empties would drag it far off; zero would leave it dead-center;
    # Word ignores space-after on a section's last paragraph, so real
    # paragraphs are the reliable spacer)
    empties = []
    nxt = target._p.getnext()
    while nxt is not None and nxt.tag == qn("w:p"):
        if ("".join(t.text or "" for t in nxt.iter(qn("w:t"))).strip()
                or nxt.find(qn("w:pPr")) is not None
                and nxt.find(qn("w:pPr")).find(qn("w:sectPr")) is not None):
            break
        empties.append(nxt)
        nxt = nxt.getnext()
    if len(empties) != 4:
        for el in empties[4:]:
            body.remove(el)
        for _ in range(4 - len(empties)):
            target._p.addnext(body.makeelement(qn("w:p"), {}))
        changed += 1

    # section break before the heading (previous sibling paragraph carrying
    # a sectPr copy WITHOUT vAlign closes the earlier, top-aligned section)
    prev = target._p.getprevious()
    has_break = (prev is not None and prev.tag == qn("w:p")
                 and prev.find(qn("w:pPr")) is not None
                 and prev.find(qn("w:pPr")).find(qn("w:sectPr")) is not None)
    if not has_break:
        brk = body.makeelement(qn("w:p"), {})
        brk_ppr = brk.makeelement(qn("w:pPr"), {})
        closing = deepcopy(tail_sectpr)
        va = closing.find(qn("w:vAlign"))
        if va is not None:
            closing.remove(va)
        brk_ppr.append(closing)
        brk.append(brk_ppr)
        target._p.addprevious(brk)
        changed += 1

    # page numbering must CONTINUE into the appendix (the closing sectPr
    # keeps the body's original pgNumType; a restart here would show
    # "Page 1" on the divider)
    pgnum = tail_sectpr.find(qn("w:pgNumType"))
    if pgnum is not None:
        tail_sectpr.remove(pgnum)
        changed += 1

    va = tail_sectpr.find(qn("w:vAlign"))
    if va is None:
        va = tail_sectpr.makeelement(qn("w:vAlign"), {})
        anchor = next((c for c in tail_sectpr
                       if c.tag in {qn(f"w:{t}") for t in _SECTPR_AFTER_VALIGN}),
                      None)
        (anchor.addprevious if anchor is not None else tail_sectpr.append)(va)
    if va.get(qn("w:val")) != "center":
        va.set(qn("w:val"), "center")
        changed += 1

    if changed:
        report.applied("Pagination",
                       "Appendix divider -> centered cover page "
                       "(horizontal + vertical)")
    return changed


def _pagination_standard(doc, report, log) -> int:
    changed = 0

    # page-break-before section headings
    for para in doc.paragraphs:
        if not para.style.name.startswith("Heading 1"):
            continue
        text = para.text.strip().lower()
        if not any(text.startswith(k) or k in text[:40]
                   for k in _BREAK_BEFORE_HEADINGS):
            continue
        ppr = para._p.get_or_add_pPr()
        if ppr.find(qn("w:pageBreakBefore")) is None:
            ppr.insert(0, ppr.makeelement(qn("w:pageBreakBefore"), {}))
            changed += 1
            report.applied("Pagination",
                           f"page break before \"{para.text.strip()[:40]}\"")

    # rows must not split across pages
    fixed_rows = 0
    for tbl in doc.tables:
        for row in tbl.rows:
            trpr = row._tr.get_or_add_trPr()
            if trpr.find(qn("w:cantSplit")) is None:
                trpr.append(trpr.makeelement(qn("w:cantSplit"), {}))
                fixed_rows += 1
    if fixed_rows:
        changed += 1
        report.applied("Pagination",
                       f"{fixed_rows} table row(s) set to never split across pages")

    # Section III label column -> 1.5" (description column absorbs the rest)
    resized = 0
    for tbl in doc.tables:
        if not _is_pastperf(tbl):
            continue
        grid = tbl._tbl.find(qn("w:tblGrid"))
        cols = grid.findall(qn("w:gridCol")) if grid is not None else []
        if len(cols) != 2:
            continue
        total = sum(int(c.get(qn("w:w"))) for c in cols)
        left = PASTPERF_LABEL_COL_TWIPS
        if int(cols[0].get(qn("w:w"))) == left:
            continue
        cols[0].set(qn("w:w"), str(left))
        cols[1].set(qn("w:w"), str(total - left))
        for row in tbl.rows:
            for idx, cell in enumerate(row.cells[:2]):
                tcpr = cell._tc.get_or_add_tcPr()
                tcw = tcpr.find(qn("w:tcW"))
                if tcw is None:
                    tcw = tcpr.makeelement(qn("w:tcW"), {})
                    tcpr.append(tcw)
                tcw.set(qn("w:type"), "dxa")
                tcw.set(qn("w:w"), str(left if idx == 0 else total - left))
        resized += 1
    if resized:
        changed += 1
        report.applied("Pagination",
                       f"{resized} past-performance table(s): label column -> "
                       f"{PASTPERF_LABEL_COL_TWIPS / 1440:.2f}\" (wider descriptions)")
    return changed


def _right_align_letterhead(p_el) -> bool:
    """Right-align a letterhead paragraph and pin its enclosing textbox (when
    it lives in one) to the right margin with no right inset -- the block's
    right edge then matches the body text's right edge exactly."""
    changed = False
    ppr = p_el.find(qn("w:pPr"))
    if ppr is None:
        ppr = p_el.makeelement(qn("w:pPr"), {})
        p_el.insert(0, ppr)
    jc = ppr.find(qn("w:jc"))
    if jc is None:
        jc = ppr.makeelement(qn("w:jc"), {})
        mark_rpr = ppr.find(qn("w:rPr"))     # rPr must stay last in pPr
        (mark_rpr.addprevious if mark_rpr is not None else ppr.append)(jc)
    if jc.get(qn("w:val")) != "right":
        jc.set(qn("w:val"), "right")
        changed = True

    anchor = next(iter(p_el.iterancestors(f"{{{_WP}}}anchor")), None)
    if anchor is not None:
        pos_h = anchor.find(f"{{{_WP}}}positionH")
        if pos_h is not None and (pos_h.get("relativeFrom") != "margin"
                                  or pos_h.find(f"{{{_WP}}}align") is None):
            pos_h.set("relativeFrom", "margin")
            for child in list(pos_h):
                pos_h.remove(child)
            align = pos_h.makeelement(f"{{{_WP}}}align", {})
            align.text = "right"
            pos_h.append(align)
            changed = True
        for bp in anchor.iter(f"{{{_WPS}}}bodyPr"):
            if bp.get("rIns") != "0":
                bp.set("rIns", "0")
                changed = True
    return changed


def proofread_document(doc, report, log=print) -> dict:
    """Enforce the document house standard in place.

    Auto-fixes font size (-> 12pt) and text color (-> black) on every table,
    normalizes the page-header letterhead (black, 9pt, right-aligned to the
    content margin), and flags border deviations (0.5pt single) for the
    human. Returns a summary dict.
    """
    sized = _normalize_font_size(doc, report, log)
    colored = _clear_colors(doc, report, log)
    letterhead = _normalize_letterhead(doc, report, log)
    paginated = _pagination_standard(doc, report, log)
    paginated += _appendix_cover(doc, report, log)
    paginated += _toc_entry_indent(doc, report, log)
    pp_borders = _apply_pastperf_borders(doc, report, log)
    borders = _flag_borders(doc, report, log)
    if sized or colored or letterhead or paginated or pp_borders or borders:
        log(f"[proofread] font-normalized {sized} table(s), color-cleared {colored}, "
            f"letterhead-fixed {letterhead} line(s), pagination-fixed {paginated}, "
            f"bordered {pp_borders} past-perf table(s), flagged {borders} for borders.")
    return {"font_tables": sized, "color_tables": colored,
            "letterhead_lines": letterhead, "pagination": paginated,
            "pp_border_tables": pp_borders, "border_flags": borders}
