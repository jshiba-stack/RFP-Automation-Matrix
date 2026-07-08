"""Rebuild a typography-damaged resume PDF onto the house template.

A resume PDF that was re-saved by a desktop PDF editor (see
``pdfutil.resume_pdf_issues``) cannot be repaired in place -- PDF is
final-form. But its *text* is intact, so this module re-typesets it:

    extract text lines (with position/size/boldness) from the bad PDF
    -> parse them into name / sections / job entries / bullets
    -> render onto the house template .docx (styles, header logo/address and
       footer come from the template; the body is written fresh)
    -> export via Word and verify no words were lost.

The result is never trusted silently: the caller flags it REVIEW so a human
proofreads the rebuilt page against the original before it ships.

House standard applied while rendering: name uppercased in the Title style,
section headers via Heading 1 (all-caps by style), company + right-tabbed
dates via Heading 2, job title via Heading 3, and the **current employer's
end year becomes "Present"** when the entry matches the firm's name.
"""

from __future__ import annotations

import copy
import math
import re
import shutil
from dataclasses import dataclass, field
from pathlib import Path

# Zones (pt, top-down): above HEADER_Y is the letterhead the template
# provides; below FOOTER_Y is the CONFIDENTIAL/filename/page footer.
_PAGE_H = 792.0
_HEADER_Y = 95.0
_FOOTER_Y = 755.0
_RIGHT_COL_X = 380.0          # runs starting right of this are the date column
_LINE_TOL = 8.0               # runs within this many pt of y = same line

_BULLET_RE = re.compile("^[\u2022\u25cf\u25aa\u00b7\ufffd]\\s*")
_FOOTER_RE = re.compile(r"^\s*(confidential\b|page\s*\d+\s*of\s*\d+)", re.I)
#: A short "Label:" lead-in (skill rows etc.) that must stay its own paragraph.
_LABEL_RE = re.compile(r"^[A-Z][\w /&()-]{0,24}:")
_DATE_RE = re.compile(r"(19|20)\d{2}|present|current", re.I)
_RANGE_RE = re.compile(r"((?:19|20)\d{2})(\s*(?:to|through|[-–—])\s*)"
                       r"((?:19|20)\d{2})", re.I)

#: Canonical section names; keys are lowercased header text as it appears.
_SECTION_ALIASES = {
    "summary": "Summary", "profile": "Summary", "professional summary": "Summary",
    "objective": "Summary",
    "technical skills": "Technical Skills", "skills": "Technical Skills",
    "technical competencies": "Technical Competencies",
    "experience": "Experience", "work experience": "Experience",
    "professional experience": "Experience", "employment history": "Experience",
    "education": "Education", "education and training": "Education",
    "certifications": "Certifications", "certification": "Certifications",
    "training": "Training",
}


@dataclass
class _Run:
    x: float
    y: float          # top-down
    size: float
    bold: bool
    text: str


@dataclass
class Line:
    y: float
    x: float
    size: float
    bold: bool
    text: str
    right_text: str = ""      # date-column portion, if any


@dataclass
class Job:
    primary: str              # the line carrying the dates (usually company)
    dates: str
    companion: str = ""       # adjacent undated line (usually job title)


@dataclass
class Parsed:
    name: str = ""
    subtitle: str = ""
    #: [(section header, items)]; item = ("para" | "bullet", text) or Job
    sections: list = field(default_factory=list)


def extract_lines(pdf_path) -> list[Line]:
    """Text lines of page 1 with layout metadata, top-to-bottom.

    Uses pypdf's extract_text visitor: each text op's combined matrix gives
    position and effective size; boldness comes from the base font name.
    """
    from pypdf import PdfReader

    runs: list[_Run] = []

    def _visit(text, cm, tm, font_dict, font_size):  # noqa: ANN001
        if not text or not text.strip():
            return
        c = tm[2] * cm[0] + tm[3] * cm[2]
        d = tm[2] * cm[1] + tm[3] * cm[3]
        e = tm[4] * cm[0] + tm[5] * cm[2] + cm[4]
        f = tm[4] * cm[1] + tm[5] * cm[3] + cm[5]
        sy = math.hypot(c, d) or 1.0
        base = str((font_dict or {}).get("/BaseFont", ""))
        runs.append(_Run(x=e, y=_PAGE_H - f, size=(font_size or 0) * sy,
                         bold="bold" in base.lower(), text=text))

    reader = PdfReader(str(pdf_path))
    page = reader.pages[0]
    page.extract_text(visitor_text=_visit)

    runs.sort(key=lambda r: (r.y, r.x))
    respace = _respacer(page.extract_text() or "")
    lines: list[Line] = []
    group: list[_Run] = []
    for r in runs:
        if group and abs(r.y - group[0].y) > _LINE_TOL:
            lines.append(_line_from(group, respace))
            group = []
        group.append(r)
    if group:
        lines.append(_line_from(group, respace))
    return lines


def _respacer(plain: str):
    """Word spacing restorer.

    A PDF-editor rewrite draws one text op per word with kerning instead of
    space glyphs, so visitor text arrives as "JordanAvery". pypdf's
    plain ``extract_text`` *does* infer those spaces (from glyph positions)
    but loses layout metadata. Marry the two: find a visitor chunk's
    space-stripped text inside the space-stripped plain text and return the
    properly spaced slice (or None when it isn't contiguous there).
    """
    compact_chars: list[str] = []
    starts: list[int] = []
    for i, ch in enumerate(plain):
        if not ch.isspace():
            compact_chars.append(ch)
            starts.append(i)
    compact = "".join(compact_chars)

    def respace(text: str) -> str | None:
        key = "".join(text.split())
        if not key:
            return ""
        pos = compact.find(key)
        if pos < 0:
            return None
        spaced = plain[starts[pos]:starts[pos + len(key) - 1] + 1]
        return re.sub(r"\s+", " ", spaced).strip()

    return respace


def _normalize(text: str) -> str:
    """Fix glyphs a PDF-editor rewrite maps to U+FFFD: a leading one is a
    bullet marker (kept for the parser), an interior one is an en dash."""
    if not text:
        return text
    head = _BULLET_RE.match(text)
    body = text[head.end():] if head else text
    body = re.sub(r"\s*�\s*", " – ", body)
    return (text[:head.end()] if head else "") + body


def _join(rs: list[_Run], respace) -> str:
    if not rs:
        return ""
    joined = respace("".join(r.text for r in rs))
    if joined is not None:
        return joined
    # not contiguous in the plain text (stream order differs): respace each
    # op alone and space-join -- editor rewrites draw one op per word anyway
    parts = [respace(r.text) or r.text for r in rs]
    return re.sub(r"\s+", " ", " ".join(parts)).strip()


def _line_from(group: list[_Run], respace) -> Line:
    """One visual line; split off a right-hand date column only when a real
    horizontal gulf separates it (editor rewrites place ordinary prose words
    past any fixed x, so absolute position alone can't decide)."""
    group = sorted(group, key=lambda r: r.x)
    split = None
    for i in range(1, len(group)):
        prev, cur = group[i - 1], group[i]
        est_end = prev.x + len(prev.text) * prev.size * 0.55
        if cur.x >= _RIGHT_COL_X and cur.x - est_end > 30:
            split = i
            break
    left = group[:split] if split else group
    right = group[split:] if split else []
    return Line(
        y=min(r.y for r in group),
        x=min(r.x for r in left),
        size=max(r.size for r in left),
        bold=any(r.bold for r in left if len(r.text.strip()) > 1),
        text=_normalize(_join(left, respace)),
        right_text=_normalize(_join(right, respace)),
    )


def _is_section_header(line: Line) -> str | None:
    """Canonical section name when the line reads as a section bar."""
    t = line.text.strip().rstrip(":").strip()
    if not t or len(t) > 40 or line.x > 120 or line.right_text:
        return None
    canon = _SECTION_ALIASES.get(t.lower())
    if canon:
        return canon
    # unknown but section-shaped: short, bold, ALL-CAPS, left margin
    if line.bold and t.isupper() and len(t.split()) <= 4:
        return t.title()
    return None


def parse_resume(lines: list[Line]) -> Parsed:
    """Structure the extracted lines. Tolerant by design: anything that isn't
    recognized as a job entry or bullet is kept as a plain paragraph, so
    content is preserved even when the layout heuristics miss."""
    body = [ln for ln in lines
            if _HEADER_Y <= ln.y <= _FOOTER_Y and ln.text
            and not _FOOTER_RE.match(ln.text)]
    header_zone = [ln for ln in lines if ln.y < _HEADER_Y and ln.text]
    out = Parsed()

    # Name: largest text in the top area (header zone or first body lines).
    top = header_zone + [ln for ln in body if ln.y < 200]
    # Letterhead address block sits at the right; the name is centered/left.
    cands = [ln for ln in top if ln.x < _RIGHT_COL_X - 60]
    if cands:
        name_ln = max(cands, key=lambda l: l.size)
        out.name = name_ln.text
        below = [l for l in cands if 0 < l.y - name_ln.y <= 25]
        if below:
            out.subtitle = min(below, key=lambda l: l.y).text
        drop = {id(name_ln)} | {id(min(below, key=lambda l: l.y))} if below else {id(name_ln)}
        body = [ln for ln in body if id(ln) not in drop]

    sections: list[tuple[str, list]] = []
    current: list = []
    sections.append(("", current))     # preamble bucket (usually empty)
    for ln in body:
        canon = _is_section_header(ln)
        if canon:
            current = []
            sections.append((canon, current))
            continue
        if ln.right_text and _DATE_RE.search(ln.right_text):
            current.append(Job(primary=ln.text, dates=ln.right_text))
            continue
        stripped = _BULLET_RE.sub("", ln.text)
        if stripped != ln.text and stripped:
            current.append(("bullet", stripped))
            continue
        # undated line right after a title-less job = its title/companion
        if (current and isinstance(current[-1], Job)
                and not current[-1].companion and not ln.bold):
            current[-1].companion = ln.text
            continue
        # wrapped prose: continue the previous paragraph unless this line
        # starts its own "Label:" row (skills tables etc.)
        if (current and isinstance(current[-1], tuple) and current[-1][0] == "para"
                and not _LABEL_RE.match(ln.text)):
            current[-1] = ("para", current[-1][1] + " " + ln.text)
            continue
        current.append(("para", ln.text))

    out.sections = [(h, items) for h, items in sections if items]
    return out


def apply_present_rule(parsed: Parsed, firm_names: list[str]) -> list[str]:
    """House standard: a current employee's entry at the firm must end in
    "Present", not a year. Returns notes describing each change made."""
    keys = [re.sub(r"[,.]?\s*(llc|inc|ltd|corp)\.?$", "", n.strip(), flags=re.I)
            for n in firm_names if n and n.strip()]
    notes = []
    for _header, items in parsed.sections:
        for it in items:
            if not isinstance(it, Job):
                continue
            if not any(k.lower() in it.primary.lower() for k in keys if k):
                continue
            new = _RANGE_RE.sub(lambda m: f"{m.group(1)} to Present", it.dates)
            if new != it.dates:
                notes.append(f"'{it.dates.strip()}' -> '{new.strip()}' ({it.primary})")
                it.dates = new
    return notes


# ---------------------------------------------------------------- rendering

def build_docx(parsed: Parsed, template_docx, out_docx) -> Path:
    """Write ``parsed`` onto a copy of the house template.

    The template supplies everything but the body: header (logo + address),
    footer (CONFIDENTIAL / FILENAME field -- Word refreshes it to the new
    file's name on export / page numbers), and the Title/Heading/List styles.
    """
    import docx
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run

    out_docx = Path(out_docx)
    shutil.copyfile(template_docx, out_docx)   # fresh mtime on purpose
    d = docx.Document(str(out_docx))
    body = d.element.body

    paras = d.paragraphs
    title_tpl = next((p._p for p in paras if p.style.name == "Title"), None)
    subtitle_tpl = paras[1]._p if len(paras) > 1 else None
    bullet_tpl = next((p._p for p in paras
                       if p.style.name == "List Paragraph"
                       and "<w:numPr>" in p._p.xml), None)
    title_tpl = copy.deepcopy(title_tpl) if title_tpl is not None else None
    subtitle_tpl = copy.deepcopy(subtitle_tpl) if subtitle_tpl is not None else None
    bullet_tpl = copy.deepcopy(bullet_tpl) if bullet_tpl is not None else None

    for el in list(body):
        if not el.tag.endswith("}sectPr"):
            body.remove(el)
    sect = body.find(qn("w:sectPr"))

    def from_template(tpl, text):
        el = copy.deepcopy(tpl)
        runs = el.findall(qn("w:r"))
        for r in runs[1:]:
            el.remove(r)
        # paragraphs cloned at the XML level must land BEFORE sectPr or Word
        # renders them after everything else
        sect.addprevious(el)
        p = Paragraph(el, d)
        if runs:
            for t in runs[0].findall(qn("w:t")):
                runs[0].remove(t)
            Run(runs[0], p).text = text
        else:
            p.add_run(text)

    def styled(style, text):
        p = d.add_paragraph(style=style)
        p.add_run().text = text            # '\t' becomes a real w:tab
        return p

    if parsed.name:
        if title_tpl is not None:
            from_template(title_tpl, parsed.name.upper())
        else:
            styled("Title", parsed.name.upper())
    if parsed.subtitle:
        if subtitle_tpl is not None:
            from_template(subtitle_tpl, parsed.subtitle)
        else:
            styled(None, parsed.subtitle)

    for header, items in parsed.sections:
        if header:
            styled("Heading 1", header)
        for it in items:
            if isinstance(it, Job):
                styled("Heading 2", f"{it.primary}\t{it.dates.strip()}")
                if it.companion:
                    styled("Heading 3", it.companion)
            elif it[0] == "bullet":
                if bullet_tpl is not None:
                    from_template(bullet_tpl, it[1])
                else:
                    styled("List Paragraph", it[1])
            else:
                styled(None, it[1])

    d.save(str(out_docx))
    return out_docx


def prepare_docx_for_conversion(src, dst) -> dict:
    """Copy a person's resume .docx to ``dst`` and apply the resume standard
    that must happen before Word converts it: hyperlink styling stripped
    (links match the surrounding text) and employment dates normalized to
    "YYYY to YYYY/Present". The original file is never modified.
    Returns ``{"links": n, "dates": n}``."""
    import docx
    from docx.oxml.ns import qn

    shutil.copyfile(src, dst)
    d = docx.Document(str(dst))
    links = strip_hyperlink_styling(d)
    dates = 0
    for t_el in d.element.body.iter(qn("w:t")):
        if not t_el.text:
            continue
        new = normalize_date_text(t_el.text)
        if new != t_el.text:
            t_el.text = new
            dates += 1
    if links or dates:
        d.save(str(dst))
    return {"links": links, "dates": dates}


#: any recognizable year range, whatever the separator
_RANGE_ANY_RE = re.compile(
    r"((?:19|20)\d{2})\s*(?:[-–—]|to|through|thru)\s*"
    r"((?:19|20)\d{2}|present|current)", re.I)
#: the house standard form
STD_RANGE_RE = re.compile(r"^(19|20)\d{2} to ((19|20)\d{2}|Present)$")


def normalize_date_text(text: str) -> str:
    """Employment-date standard: every year range reads "YYYY to YYYY" or
    "YYYY to Present" (hyphens/dashes/'Current' normalized)."""
    def _rep(m):
        end = m.group(2)
        if end.lower() in ("present", "current"):
            end = "Present"
        return f"{m.group(1)} to {end}"
    return _RANGE_ANY_RE.sub(_rep, text)


#: Word's default hyperlink character color
HYPERLINK_BLUE = "0563C1"


def strip_hyperlink_styling(doc) -> int:
    """Resume standard: links render like the text around them.

    Removes hyperlink character styling (Hyperlink run style, the Word link
    blue, underline) from every body run that is inside a ``w:hyperlink``,
    carries a Hyperlink-ish run style, or is colored the default link blue --
    the run then inherits its paragraph's color. Returns runs changed.
    """
    from docx.oxml.ns import qn

    n = 0
    for r_el in doc.element.body.iter(qn("w:r")):
        rpr = r_el.find(qn("w:rPr"))
        if rpr is None:
            continue
        rstyle = rpr.find(qn("w:rStyle"))
        linkish_style = (rstyle is not None
                         and "hyperlink" in (rstyle.get(qn("w:val")) or "").lower())
        color = rpr.find(qn("w:color"))
        link_blue = (color is not None
                     and (color.get(qn("w:val")) or "").upper() == HYPERLINK_BLUE)
        in_link = r_el.getparent().tag == qn("w:hyperlink")
        if not (in_link or linkish_style or link_blue):
            continue
        changed = False
        if linkish_style:
            rpr.remove(rstyle)
            changed = True
        if color is not None:
            rpr.remove(color)
            changed = True
        u = rpr.find(qn("w:u"))
        if u is not None:
            rpr.remove(u)
            changed = True
        if changed:
            n += 1
    return n


# ------------------------------------------------------------- verification

_WORD_RE = re.compile(r"[A-Za-z0-9][A-Za-z0-9./&+#'-]*")


def _words(text: str) -> set[str]:
    return {w.lower() for w in _WORD_RE.findall(text)}


def missing_words(source_pdf, rebuilt_pdf, ignore: set[str] | None = None) -> set[str]:
    """Body words of the source that did not survive into the rebuild.

    Header/footer zones are excluded from the source side (the template
    replaces them by design). ``ignore`` covers deliberate edits, e.g. the
    end year the Present rule removed. Matching is done against the rebuilt
    text with ALL whitespace stripped, because pypdf's extraction sometimes
    splits a word at a line-break kern ("performan ce").
    """
    from pypdf import PdfReader

    src = set()
    for ln in extract_lines(source_pdf):
        if _HEADER_Y <= ln.y <= _FOOTER_Y:
            src |= _words(ln.text) | _words(ln.right_text)
    reader = PdfReader(str(rebuilt_pdf))
    got = "".join("".join((pg.extract_text() or "").split()).lower()
                  for pg in reader.pages)
    return {w for w in src if w not in got} - (ignore or set())


# -------------------------------------------------------------- entry point

def rebuild(pdf_path, template_docx, out_dir, firm_names=()) -> dict:
    """Full pipeline: parse the damaged PDF, render, Word-export, verify.

    Returns ``{ok, docx, pdf, notes, error}``. ``ok`` is True only when Word
    produced a PDF and no body words were lost. Results are cached in
    ``out_dir``: an existing rebuild newer than its source is reused.
    """
    from . import pdfutil

    pdf_path = Path(pdf_path)
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    out_docx = out_dir / f"{pdf_path.stem} (REBUILT).docx"
    out_pdf = out_docx.with_suffix(".pdf")

    try:
        src_mtime = pdf_path.stat().st_mtime
        if (out_pdf.exists() and out_docx.exists()
                and out_pdf.stat().st_mtime > src_mtime):
            return {"ok": True, "docx": out_docx, "pdf": out_pdf,
                    "notes": ["reused cached rebuild"], "error": ""}

        parsed = parse_resume(extract_lines(pdf_path))
        if not parsed.sections:
            return {"ok": False, "error": "couldn't parse any sections",
                    "docx": None, "pdf": None, "notes": []}
        for _h, items in parsed.sections:   # date-format standard
            for it in items:
                if isinstance(it, Job):
                    it.dates = normalize_date_text(it.dates)
        notes = apply_present_rule(parsed, list(firm_names))
        # years the Present rule dropped shouldn't count as lost words
        dropped = {m.group(3).lower()
                   for n in notes for m in [_RANGE_RE.search(n)] if m}
        build_docx(parsed, template_docx, out_docx)
        exported = pdfutil.export_pdf(out_docx, out_pdf)
        if not exported:
            return {"ok": False, "error": "Word export failed (Word unavailable?)",
                    "docx": out_docx, "pdf": None, "notes": notes}
        lost = missing_words(pdf_path, out_pdf, ignore=dropped | {"present"})
        if lost:
            sample = ", ".join(sorted(lost)[:8])
            return {"ok": False, "docx": out_docx, "pdf": out_pdf, "notes": notes,
                    "error": f"rebuild lost words ({sample}) -- keeping the original"}
        return {"ok": True, "docx": out_docx, "pdf": out_pdf, "notes": notes,
                "error": ""}
    except Exception as exc:  # noqa: BLE001 - a broken source must never kill the build
        return {"ok": False, "error": f"rebuild failed: {exc}",
                "docx": None, "pdf": None, "notes": []}
