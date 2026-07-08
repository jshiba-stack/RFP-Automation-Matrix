"""PDF measurement helpers for the compliance checklist.

We never render the deliverable PDF ourselves -- only Word produces a faithful
submittal. These helpers MEASURE a PDF the user exported (size, page count), and
optionally drive Word (via docx2pdf, Windows only) to produce a throwaway PDF
purely for measurement when ``auto_export_pdf`` is enabled.
"""

from __future__ import annotations

import math
import os
import re
from pathlib import Path

#: PDF base-14 fonts every viewer ships; embedding is never required for these.
_STANDARD_FONTS = {
    "Helvetica", "Helvetica-Bold", "Helvetica-Oblique", "Helvetica-BoldOblique",
    "Courier", "Courier-Bold", "Courier-Oblique", "Courier-BoldOblique",
    "Times-Roman", "Times-Bold", "Times-Italic", "Times-BoldItalic",
    "Symbol", "ZapfDingbats",
}

_LETTER = (612.0, 792.0)  # points


def pdf_size_mb(path) -> float:
    return os.path.getsize(path) / (1024 * 1024)


def pdf_page_count(path) -> int | None:
    try:
        from pypdf import PdfReader
    except Exception:
        return None
    try:
        return len(PdfReader(str(path)).pages)
    except Exception:
        return None


def find_companion_pdf(docx_path) -> Path | None:
    """A PDF the user exported next to the .docx (same stem) if one exists."""
    p = Path(docx_path)
    cand = p.with_suffix(".pdf")
    return cand if cand.exists() else None


def export_pdf(docx_path, out_pdf=None) -> Path | None:
    """Export docx -> pdf via Word (docx2pdf). Returns the path or None.

    Windows + Word only. Used to produce the submittal body PDF (and to
    convert a .docx resume) during assembly; safe to call from a worker
    thread (COM is initialized for the calling thread).
    """
    try:
        from docx2pdf import convert
    except Exception:
        return None
    try:  # COM must be initialized per-thread (the dashboard uses a worker)
        import pythoncom
        pythoncom.CoInitialize()
    except Exception:  # noqa: BLE001
        pass
    out = Path(out_pdf) if out_pdf else Path(docx_path).with_suffix(".measure.pdf")
    try:
        convert(str(docx_path), str(out))
    except Exception:
        return None
    return out if out.exists() else None


def pdf_editor_rewrite(path) -> str | None:
    """Name of the PDF editor that re-saved a Word-exported PDF, or None.

    Word writes the same app name to both Creator and Producer; a desktop PDF
    editor (PDFescape et al.) replaces Creator but keeps Word's Producer.
    Those editors re-write the text layer and routinely mangle fonts and
    glyph scaling, so a re-saved file is a worse pick than a straight export.
    """
    try:
        from pypdf import PdfReader
        meta = PdfReader(str(path)).metadata or {}
    except Exception:
        return None
    creator = str(meta.get("/Creator") or "").strip()
    producer = str(meta.get("/Producer") or "")
    if creator and "word" in producer.lower() and "word" not in creator.lower():
        return creator
    return None


def _font_descriptor(font_obj):
    desc = font_obj.get("/FontDescriptor")
    if desc is None and "/DescendantFonts" in font_obj:
        desc = font_obj["/DescendantFonts"][0].get_object().get("/FontDescriptor")
    return desc.get_object() if desc is not None else None


def resume_pdf_issues(path) -> list[str]:
    """Typography defects that make a resume page look wrong in the submittal.

    Each returned string reads as a verb phrase after "<name>'s resume PDF ..."
    ([] = clean). Checks: re-saved by a PDF editor, text non-uniformly scaled
    (the stretched-words look), non-embedded fonts (viewers substitute and
    distort), and off-Letter page size.
    """
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
    except Exception:
        return []  # unreadable here -> the merge itself will surface it
    issues: list[str] = []

    editor = pdf_editor_rewrite(path)
    if editor:
        issues.append(f"was re-saved with {editor} after the Word export "
                      "(text layer re-written)")

    # Text stretched: glyphs drawn with unequal x/y scale (Tm/cm transforms).
    stretched: list[float] = []
    def _visit(text, cm, tm, font_dict, font_size):  # noqa: ANN001
        if not text or len(text.strip()) < 3 or not font_size:
            return
        a = tm[0] * cm[0] + tm[1] * cm[2]
        b = tm[0] * cm[1] + tm[1] * cm[3]
        c = tm[2] * cm[0] + tm[3] * cm[2]
        d = tm[2] * cm[1] + tm[3] * cm[3]
        sx, sy = math.hypot(a, b), math.hypot(c, d)
        if sx and sy and abs(sy / sx - 1) > 0.05:
            stretched.append(sy / sx)

    bad_size = None
    nonembedded: set[str] = set()
    for page in reader.pages:
        mb = page.mediabox
        w, h = float(mb.width), float(mb.height)
        if bad_size is None and (abs(w - _LETTER[0]) > 3 or abs(h - _LETTER[1]) > 3):
            bad_size = (w, h)
        try:
            fonts = (page.get("/Resources") or {}).get("/Font") or {}
            for ref in fonts.values():
                font = ref.get_object()
                base = str(font.get("/BaseFont") or "").lstrip("/")
                name = base.split("+", 1)[-1]  # drop the ABCDEF+ subset prefix
                if not name or name in _STANDARD_FONTS:
                    continue
                desc = _font_descriptor(font)
                if desc is None or not any(
                        k in desc for k in ("/FontFile", "/FontFile2", "/FontFile3")):
                    nonembedded.add(name)
        except Exception:  # noqa: BLE001 - odd resource tree -> skip font check
            pass
        try:
            page.extract_text(visitor_text=_visit)
        except Exception:  # noqa: BLE001 - extraction failure -> skip stretch check
            pass

    if len(stretched) >= 5:
        worst = max(stretched, key=lambda r: abs(r - 1))
        pct = round(abs(worst - 1) * 100)
        how = "taller" if worst > 1 else "flatter"
        issues.append(f"draws its text non-uniformly scaled (up to {pct}% {how} "
                      "than designed) -- words render stretched")
    if nonembedded:
        issues.append("uses fonts that aren't embedded "
                      f"({', '.join(sorted(nonembedded))}) -- viewers substitute "
                      "and distort them")
    if bad_size:
        issues.append(f"has a {bad_size[0]:g} x {bad_size[1]:g} pt page instead "
                      "of Letter (612 x 792)")
    return issues


def merge_pdfs(paths) -> bytes:
    """Concatenate PDFs (in order) into one; returns the merged bytes."""
    from io import BytesIO

    from pypdf import PdfWriter

    writer = PdfWriter()
    for p in paths:
        writer.append(str(p))
    bio = BytesIO()
    writer.write(bio)
    return bio.getvalue()


# --- letterhead standardization ----------------------------------------------
# The letterhead block (firm name + address + site, top-right of every page)
# drifts across resume files: different sizes, positions, colors, address
# spellings. The submittal BODY is the house anchor: its letterhead is
# extracted, re-set via Word as a transparent one-page "stamp", and every
# resume page gets its old block whited out and the stamp merged on top --
# so the whole deliverable shows one identical letterhead.

#: top-right page region a letterhead lives in (top-down pt, Letter page)
_LH_X0, _LH_Y1 = 300.0, 95.0
#: whiteout rectangle (bottom-up PDF coords): x, y, w, h
_LH_WHITE = (350.0, 792.0 - 92.0, 262.0, 72.0)


def _zone_runs(pdf_path, page_index=0) -> list[tuple[float, float, float, str]]:
    """(y_topdown, x, size, text) for text in the letterhead zone."""
    import math

    from pypdf import PdfReader

    out = []

    def _visit(text, cm, tm, font_dict, font_size):  # noqa: ANN001
        if not text or not text.strip():
            return
        c = tm[2] * cm[0] + tm[3] * cm[2]
        d = tm[2] * cm[1] + tm[3] * cm[3]
        e = tm[4] * cm[0] + tm[5] * cm[2] + cm[4]
        f = tm[4] * cm[1] + tm[5] * cm[3] + cm[5]
        y = 792.0 - f
        if e >= _LH_X0 and y <= _LH_Y1:
            out.append((y, e, (font_size or 0) * (math.hypot(c, d) or 1.0),
                        text))

    try:
        reader = PdfReader(str(pdf_path))
        reader.pages[page_index].extract_text(visitor_text=_visit)
    except Exception:  # noqa: BLE001
        return []
    return sorted(out)


def _zone_lines(runs) -> list[tuple[float, float, float, str]]:
    """Group zone runs into (baseline_y, x, size, text) lines (2pt tol)."""
    lines: list[list] = []
    for y, x, size, text in runs:
        if lines and abs(y - lines[-1][0][0]) <= 2.0:
            lines[-1].append((y, x, size, text))
        else:
            lines.append([(y, x, size, text)])
    out = []
    for grp in lines:
        grp.sort(key=lambda r: r[1])
        # concatenate raw run text (Word ops carry their own spaces; joining
        # with " " would break tokens pypdf splits, e.g. "tdc"+"-"+"hawaii")
        text = re.sub(r"\s+", " ", "".join(r[3] for r in grp)).strip()
        out.append((grp[0][0], grp[0][1], max(r[2] for r in grp), text))
    return out


def logo_top(pdf_path, page_index=0) -> float | None:
    """Top edge (top-down pt) of the letterhead logo image on a page, or
    None. Watches image-XObject Do operators; the placement CTM gives the
    drawn rectangle. Used to align the letterhead stamp with each page's
    logo the way the body aligns its own."""
    from pypdf import PdfReader

    try:
        reader = PdfReader(str(pdf_path))
        page = reader.pages[page_index]
        xobjs = (page.get("/Resources") or {}).get("/XObject")
        xobjs = xobjs.get_object() if xobjs is not None else {}
    except Exception:  # noqa: BLE001
        return None
    tops: list[float] = []

    def _before(op, operands, cm, tm):  # noqa: ANN001
        if op != b"Do" or not operands:
            return
        try:
            xo = xobjs.get(str(operands[0]))
            if xo is None or xo.get_object().get("/Subtype") != "/Image":
                return
        except Exception:  # noqa: BLE001
            return
        w, h, x, y = cm[0], cm[3], cm[4], cm[5]
        top = 792.0 - (y + h)
        if top < 150 and x < 300 and w > 30 and h > 10:  # top-left, logo-sized
            tops.append(top)

    try:
        page.extract_text(visitor_operand_before=_before)
    except Exception:  # noqa: BLE001
        return None
    return min(tops) if tops else None


def letterhead_spec(body_pdf) -> dict | None:
    """The house letterhead as measured from the body PDF: text lines plus
    geometry (first baseline, line pitch, font size, logo top edge)."""
    for page_index in range(3):
        lines = _zone_lines(_zone_runs(body_pdf, page_index))
        if len(lines) >= 2:
            ys = [l[0] for l in lines]
            pitch = (ys[-1] - ys[0]) / (len(ys) - 1)
            return {"lines": [l[3] for l in lines],
                    "baseline": ys[0], "pitch": pitch,
                    "size": round(max(l[2] for l in lines) * 2) / 2,
                    "logo_top": logo_top(body_pdf, page_index)}
    return None


def build_letterhead_stamp(body_pdf, out_pdf) -> Path | None:
    """A one-page PDF holding ONLY the body's letterhead block, position-
    calibrated to match the body within ~0.7pt. Cached next to ``out_pdf``
    (reused while the body's letterhead text is unchanged). None when the
    body has no letterhead or Word is unavailable.
    """
    import json

    spec = letterhead_spec(body_pdf)
    if spec is None:
        return None
    out_pdf = Path(out_pdf)
    sidecar = out_pdf.with_suffix(".json")
    cache_key = {"lines": spec["lines"], "size": spec["size"]}
    try:
        if (out_pdf.exists()
                and json.loads(sidecar.read_text(encoding="utf-8")) == cache_key):
            return out_pdf
    except Exception:  # noqa: BLE001 - stale/absent sidecar -> rebuild
        pass

    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.shared import Pt, RGBColor

    out_pdf.parent.mkdir(parents=True, exist_ok=True)
    tmp_docx = out_pdf.with_suffix(".docx")
    top = spec["baseline"] - spec["size"]  # first guess

    def _write(top_pt):
        d = docx.Document()
        sec = d.sections[0]
        sec.page_width, sec.page_height = Pt(612), Pt(792)
        # right-aligned block flush with the content margin: the standard
        # puts the block's right edge exactly at the body text's right edge
        sec.left_margin, sec.right_margin = Pt(72), Pt(72)
        sec.top_margin, sec.bottom_margin = Pt(max(top_pt, 0)), Pt(18)
        sec.header_distance = Pt(0)
        for line in spec["lines"]:
            p = d.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            pf = p.paragraph_format
            pf.space_before = pf.space_after = Pt(0)
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(spec["pitch"])
            r = p.add_run(line.strip())
            r.font.name = "Calibri"
            r.font.size = Pt(spec["size"])
            r.font.color.rgb = RGBColor(0, 0, 0)
        d.save(str(tmp_docx))
        return export_pdf(tmp_docx, out_pdf)

    try:
        for _ in range(3):  # calibrate the vertical against our own output
            if _write(top) is None:
                return None
            got = _zone_lines(_zone_runs(out_pdf))
            if not got:
                return None
            dy = spec["baseline"] - got[0][0]
            if abs(dy) <= 0.7:
                break
            top += dy
        sidecar.write_text(json.dumps(cache_key), encoding="utf-8")
        return out_pdf
    finally:
        tmp_docx.unlink(missing_ok=True)


def _whiteout_pdf_bytes() -> bytes:
    """Minimal one-page PDF that paints a white rectangle over the
    letterhead zone (merged onto a page, it hides the old block)."""
    x, y, w, h = _LH_WHITE
    content = f"q 1 1 1 rg {x:g} {y:g} {w:g} {h:g} re f Q".encode()
    objs = {
        1: b"<< /Type /Catalog /Pages 2 0 R >>",
        2: b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        3: b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
           b"/Contents 4 0 R /Resources << >> >>",
        4: b"<< /Length %d >>\nstream\n%s\nendstream" % (len(content), content),
    }
    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for num in sorted(objs):
        offsets[num] = len(out)
        out += b"%d 0 obj\n" % num + objs[num] + b"\nendobj\n"
    xref = len(out)
    out += b"xref\n0 5\n0000000000 65535 f \n"
    for num in range(1, 5):
        out += b"%010d 00000 n \n" % offsets[num]
    out += b"trailer\n<< /Size 5 /Root 1 0 R >>\nstartxref\n%d\n%%%%EOF\n" % xref
    return bytes(out)


def stamp_letterhead(resume_pdf, stamp_pdf, out_path, dy: float = 0.0):
    """Replace page 1's letterhead block with the house stamp.

    ``dy`` shifts the stamp down (top-down pt) so the block sits relative to
    THIS page's logo the way the body's block sits relative to its own
    (resume logos render lower than the body's).

    Safety first: every text line found in the letterhead zone must read as
    letterhead (pattern match, or close to one of the stamp's own lines) --
    otherwise nothing is touched and the reason is returned.
    Returns ``(stamped_path, "")`` or ``(None, reason)``.
    """
    from difflib import SequenceMatcher
    from io import BytesIO

    from pypdf import PdfReader, PdfWriter, Transformation

    from .proofread import _LETTERHEADISH

    spec_lines = [l[3] for l in _zone_lines(_zone_runs(stamp_pdf))]
    zone = _zone_lines(_zone_runs(resume_pdf))
    for _y, _x, _size, text in zone:
        if _LETTERHEADISH.search(text):
            continue
        if any(SequenceMatcher(None, text.lower(), s.lower()).ratio() > 0.5
               for s in spec_lines):
            continue
        return None, f"unexpected text in the header zone: \"{text[:40]}\""

    try:
        white = PdfReader(BytesIO(_whiteout_pdf_bytes())).pages[0]
        stamp = PdfReader(str(stamp_pdf)).pages[0]
        reader = PdfReader(str(resume_pdf))
        writer = PdfWriter()
        for i, page in enumerate(reader.pages):
            if i == 0:
                page.merge_page(white)
                if dy:  # visually down = negative y in PDF coords
                    page.merge_transformed_page(
                        stamp, Transformation().translate(0, -dy))
                else:
                    page.merge_page(stamp)
            writer.add_page(page)
        out_path = Path(out_path)
        with open(out_path, "wb") as fh:
            writer.write(fh)
        return out_path, ""
    except Exception as exc:  # noqa: BLE001 - never let stamping kill a build
        return None, f"stamping failed: {exc}"
