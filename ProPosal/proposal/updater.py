"""PRIMARY mode: smart copy-and-update of a previous submittal.

Opens the base .docx, auto-applies the mechanical edits, then syncs the
store-managed content into the document. The base file is never mutated -- we
save a new draft.

Auto-applied (each still reported):
  * Fiscal year -- from the store / an explicit override (the document's own
    year is kept when neither says otherwise)
  * Cover + letter date -> the build date
  * Ongoing Capacity end-dates ('2025+') -> '<as-of-year>+'
  * Store sync: NEW Section II rows / III blocks / IV rows are appended by
    cloning the document's own formatting; CHANGED fields on matched entries
    are updated in place. Document content absent from the store is left alone.

Flagged (never silently changed):
  * A match whose runs have mixed formatting (UNSAFE EDIT)
  * A section table that can't be found / has no block to clone (MISSING)
"""

from __future__ import annotations

import datetime as _dt
import re

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell

from . import docx_map, skills
from .docx_edit import (append_cloned_row, para_text, rebuild_table_body,
                        replace_in_paragraph, set_cell_lines, set_cell_text)
from .flags import KIND_ADD, KIND_MISSING, KIND_UNSAFE, Report


def _coerce_date(value, fallback: _dt.date) -> _dt.date:
    if value in (None, ""):
        return fallback
    if isinstance(value, _dt.datetime):
        return value.date()
    if isinstance(value, _dt.date):
        return value
    s = str(value).strip()
    # ISO 'YYYY-MM-DD'
    try:
        return _dt.date.fromisoformat(s)
    except ValueError:
        pass
    # 'Month D, YYYY'
    for fmt in ("%B %d, %Y", "%b %d, %Y"):
        try:
            return _dt.datetime.strptime(s, fmt).date()
        except ValueError:
            continue
    raise ValueError(f"Unrecognized cover_date: {value!r}")


def _fmt_date(d: _dt.date) -> str:
    return f"{d:%B} {d.day}, {d.year}"


def build(
    base_path,
    store: dict | None = None,
    *,
    target_fy: int | None = None,
    cover_date=None,
    today: _dt.date | None = None,
    resumes_dir=None,
    log=print,
):
    """Return ``(Document, Report)``. Caller saves the document."""
    store = store or {}
    opp = store.get("opportunity", {})
    today = today or _dt.date.today()

    doc = Document(str(base_path))
    report = Report(base=str(base_path))

    detected_fy = docx_map.detect_fiscal_year(doc)
    if target_fy is None:
        # No bump by default: the document keeps its own year unless the store
        # (or an explicit override) says otherwise.
        target_fy = opp.get("fiscal_year") or detected_fy
    if target_fy is None:
        report.flag("document", "Could not detect a fiscal year in the document.", KIND_MISSING)

    cover = _coerce_date(cover_date if cover_date is not None else opp.get("cover_date"), today)
    as_of_year = cover.year

    log(f"Detected FY {detected_fy} -> target FY {target_fy}; cover date {_fmt_date(cover)}")

    if target_fy and target_fy != detected_fy:
        apply_fiscal_year(doc, target_fy, report)
    apply_cover_dates(doc, cover, report)
    apply_ongoing_end_dates(doc, as_of_year, report)
    apply_store_sync(doc, store, as_of_year, report)
    if resumes_dir:
        from . import resumes
        resumes.add_resume_flags(report, store, resumes_dir)

    return doc, report


# --- reusable scalar edits (shared with generate mode) ---------------------

def apply_fiscal_year(doc, target_fy: int, report: Report) -> None:
    """Bump every fiscal-year reference: title, inline letter text, footer FY##."""
    for p, kind in docx_map.find_fiscal_year_paragraphs(doc):
        if kind == "fiscal_year":
            res = replace_in_paragraph(p, r"Fiscal Year 20\d\d", f"Fiscal Year {target_fy}")
        else:  # inline 'for the fiscal year YYYY'
            res = replace_in_paragraph(
                p, r"(?<=fiscal year )20\d\d", str(target_fy), flags=re.IGNORECASE
            )
        _record(report, "Fiscal year", res)

    fy2 = f"FY{target_fy % 100:02d}"  # surgical: leaves the footer PAGE field intact
    for p, si in docx_map.footer_paragraphs(doc):
        if docx_map.RE_FOOTER_FY.search(para_text(p)):
            res = replace_in_paragraph(p, r"FY\d{2}", fy2)
            _record(report, f"Footer (section {si})", res)


def apply_cover_dates(doc, cover: _dt.date, report: Report) -> None:
    new_date = _fmt_date(cover)
    date_paras = docx_map.find_cover_date_paragraphs(doc)
    if not date_paras:
        report.flag("cover", "No cover/letter date paragraph found.", KIND_MISSING)
    for p, label in date_paras:
        res = replace_in_paragraph(p, docx_map.RE_FULL_DATE, new_date)
        _record(report, label, res)


def apply_ongoing_end_dates(doc, as_of_year: int, report: Report) -> None:
    """Refresh ongoing Capacity end-dates ('2025+') to '<as_of_year>+'."""
    tbl = docx_map.find_capacity_table(doc)
    if tbl is None:
        report.flag(
            "Capacity",
            "Capacity table not found (header Client/Project/Start Date/End Date).",
            KIND_MISSING,
        )
        return
    header = [c.text.strip().lower() for c in tbl.rows[0].cells]
    try:
        end_col = header.index("end date")
    except ValueError:
        end_col = len(header) - 1
    for ri, row in enumerate(tbl.rows[1:], start=1):
        cell = row.cells[end_col]
        for para in cell.paragraphs:
            if docx_map.RE_ONGOING_YEAR.search(para_text(para)):
                res = replace_in_paragraph(para, docx_map.RE_ONGOING_YEAR, str(as_of_year))
                _record(report, f"Capacity r{ri} End Date", res, what="ongoing end-date")


def _record(report: Report, location: str, res, what: str = "value") -> None:
    if not res.matched:
        return
    if res.applied:
        report.applied(location, f"updated {what}", res.old, res.new)
    elif not res.rpr_uniform:
        report.flag(
            location,
            f"{what} spans mixed formatting; left unchanged -- edit by hand.",
            KIND_UNSAFE,
            res.old,
            res.new,
        )


def _norm_name(name: str) -> str:
    """Normalize for matching: case, commas, and any whitespace (incl. newlines).

    The doc's block headers wrap the client over two lines with no comma, while
    a store entry naturally reads 'City & County of Honolulu, Department of...'
    -- both normalize to the same key.
    """
    return re.sub(r"[\s,]+", " ", name).strip().lower()


# --- store sync: append/update Sections II, III, IV in the base document ----

def _end_str(project: dict, as_of_year: int) -> str:
    end = project.get("end")
    if end in (None, "", "ongoing"):
        return f"{as_of_year}+"
    return str(end)


def _update_cell(cell, new_text: str, report: Report, location: str, what: str) -> None:
    """Set a cell's text when the store value differs (empty store value = skip)."""
    if not new_text:
        return
    old = cell.text.strip()
    if _norm_name(old) == _norm_name(new_text):
        return
    set_cell_text(cell, new_text)
    report.applied(location, f"updated {what} from store", old, new_text)


def sync_past_performance(doc, past_performance: list, report: Report) -> None:
    """Public entry for Section III sync (shared with generate mode)."""
    _sync_past_performance(doc, past_performance, report)


def apply_store_sync(doc, store: dict, as_of_year: int, report: Report) -> None:
    """Sync store-managed content into the document (PRIMARY mode).

    New entries are appended by cloning the section's own formatting; changed
    fields on matched entries are updated in place. Rows/blocks in the document
    that the store doesn't mention are left untouched (delete by hand).
    """
    _sync_categories(doc, store.get("categories") or [], report)
    _sync_qualifications(doc, store.get("personnel") or [], report)
    _sync_past_performance(doc, store.get("past_performance") or [], report)
    _sync_capacity(doc, store.get("projects") or [], as_of_year, report)


# tcPr children that must follow <w:vMerge> in a schema-valid cell.
_TCPR_AFTER_VMERGE = [qn(f"w:{t}") for t in (
    "tcBorders", "shd", "noWrap", "tcMar", "textDirection", "tcFitText",
    "vAlign", "hideMark")]


def _set_tc_vmerge(tc, restart: bool) -> None:
    """Set a <w:tc>'s vertical-merge state: restart (top of span) or continue."""
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:vMerge")):
        tcPr.remove(old)
    vm = tcPr.makeelement(qn("w:vMerge"), {})
    if restart:
        vm.set(qn("w:val"), "restart")
    anchor = next((c for c in tcPr if c.tag in _TCPR_AFTER_VMERGE), None)
    (anchor.addprevious if anchor is not None else tcPr.append)(vm)


def _vmerge_first_column(tbl, start_row: int, end_row: int, text: str) -> None:
    """Vertically merge column 0 of table rows [start_row, end_row] into one cell
    showing ``text`` (rows 0-based; row 0 is the header).

    Works on the raw <w:tc> elements: once a cell is ``vMerge=restart``, python-docx
    returns the *origin* cell for lower rows' ``cells[0]``, so touching cells via the
    grid would keep hitting the top cell.
    """
    trs = tbl._tbl.tr_lst
    for offset, ri in enumerate(range(start_row, end_row + 1)):
        tc = trs[ri].tc_lst[0]
        set_cell_text(_Cell(tc, tbl), text if offset == 0 else "")
        _set_tc_vmerge(tc, restart=(offset == 0))


def _clear_table_font_color(tbl) -> None:
    """Drop explicit run colors in a table so text uses the document default.

    Removes leftover red annotation coloring (the firm marks in-progress rows red);
    a submittal table should be the standard black.
    """
    for row in tbl.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    rpr = run._r.find(qn("w:rPr"))
                    if rpr is not None:
                        col = rpr.find(qn("w:color"))
                        if col is not None:
                            rpr.remove(col)


def _sync_categories(doc, categories: list, report: Report) -> None:
    """Section I: rebuild the Categories table as a clean, FY-standard block.

    Rather than patch rows in place (which let stale letters and red annotation
    color survive), the table is rebuilt from ``skills.finalize_categories``:
    letters reconciled against the FY taxonomy, sorted a-x, duplicates combined,
    column 2 set to the canonical category name, and font color cleared. This is
    idempotent and independent of accumulated store state.
    """
    if not categories:
        return
    tables = docx_map.find_table_by_signature(doc, docx_map.SIG_CATEGORIES)
    if not tables:
        report.flag("Categories", "Categories (DIT #) table not found; add the store "
                    "categories manually.", KIND_MISSING)
        return
    finalized = skills.finalize_categories(categories)
    if not finalized:
        return
    n = write_finalized_categories(tables[0], finalized)
    report.applied(
        "Categories",
        f"rebuilt Section I from store: uppercase A-X, sorted, canonical names "
        f"(X keeps per-skill titles), merged catch-all, cleared colors ({n} rows)")


def write_finalized_categories(tbl, finalized: list[dict]) -> int:
    """Rebuild the Categories table from a finalized list (shared by both engines).

    Non-X rows use the canonical FY name; the catch-all X block keeps each
    specialty's own title/description and shares a single merged "X" cell; all
    explicit font colors are cleared. Returns the number of rows written.
    """
    # rebuild letter + name; the description is set separately so its per-item
    # line breaks survive (rebuild_table_body collapses to a single run).
    rows = [[r["dit_number"], r["name"], ""] for r in finalized]
    n = rebuild_table_body(tbl, rows)
    for i, r in enumerate(finalized):
        set_cell_lines(tbl.rows[i + 1].cells[2], r["description"])
    x_idx = [i for i, r in enumerate(finalized) if r.get("catchall")]
    if len(x_idx) >= 2:
        _vmerge_first_column(tbl, x_idx[0] + 1, x_idx[-1] + 1, finalized[x_idx[0]]["dit_number"])
    _clear_table_font_color(tbl)
    return n


def _sync_capacity(doc, projects: list, as_of_year: int, report: Report) -> None:
    if not projects:
        return
    tbl = docx_map.find_capacity_table(doc)
    if tbl is None:
        report.flag("Capacity", "Capacity table not found; add the store projects manually.",
                    KIND_MISSING)
        return
    rows_by_key = {}
    for row in tbl.rows[1:]:
        key = (_norm_name(row.cells[0].text), _norm_name(row.cells[1].text))
        rows_by_key.setdefault(key, row)
    for proj in projects:
        client = str(proj.get("client", ""))
        project = str(proj.get("project", ""))
        key = (_norm_name(client), _norm_name(project))
        start = str(proj.get("start_year", "") or "")
        end = _end_str(proj, as_of_year)
        row = rows_by_key.get(key)
        if row is None:
            append_cloned_row(tbl, [client, project, start, end])
            report.applied("Capacity", "added project row from store",
                           new=f"{client} / {project} ({start}–{end})")
        else:
            loc = f"Capacity '{client or project}'"
            _update_cell(row.cells[2], start, report, loc, "Start Date")
            _update_cell(row.cells[3], end, report, loc, "End Date")


def _sync_qualifications(doc, personnel: list, report: Report) -> None:
    if not personnel:
        return
    tables = docx_map.find_table_by_signature(doc, docx_map.SIG_QUALIFICATIONS)
    if not tables:
        report.flag("Qualifications", "Resource/Qualifications table not found; "
                    "add the store personnel manually.", KIND_MISSING)
        return
    tbl = tables[0]
    rows_by_name = {}
    for row in tbl.rows[1:]:
        rows_by_name.setdefault(_norm_name(row.cells[0].text), row)
    for person in personnel:
        name = str(person.get("name", "")).strip()
        if not name:
            continue
        quals = str(person.get("qualifications") or person.get("role") or "")
        row = rows_by_name.get(_norm_name(name))
        if row is None:
            append_cloned_row(tbl, [name, quals])
            report.applied("Qualifications", "added resource row from store", new=name)
        else:
            _update_cell(row.cells[1], quals, report,
                         f"Qualifications '{name}'", "qualifications")


# Past-performance block row labels -> store fields.
_PP_FIELD_BY_LABEL = {
    "client": "client",
    "project": "project",
    "client contact": "contact",
    "client phone": "phone",
    "detailed scope of work": "scope",
    "issue resolution": "issue_resolution",
}


def _pp_blocks(doc) -> list[tuple[Table, str, str]]:
    """Every 2-col past-performance block as (table, norm_client, norm_project)."""
    out = []
    for t in doc.tables:
        if not t.rows or len(t.rows[0].cells) != 2:
            continue
        if t.rows[0].cells[0].text.strip().lower() != docx_map.SIG_PASTPERF_FIRST_CELL:
            continue
        client = t.rows[0].cells[1].text
        project = t.rows[1].cells[1].text if len(t.rows) > 1 else ""
        out.append((t, _norm_name(client), _norm_name(project)))
    return out


def _fill_pp_block(tbl: Table, entry: dict, report: Report, location: str,
                   *, update_only: bool) -> None:
    """Write store fields into a block's value cells (matched by row label)."""
    for row in tbl.rows:
        field = _PP_FIELD_BY_LABEL.get(_norm_name(row.cells[0].text))
        if field is None:
            continue
        value = str(entry.get(field, "") or "")
        if update_only:
            _update_cell(row.cells[1], value, report, location, field.replace("_", " "))
        elif value:
            set_cell_text(row.cells[1], value)


def _pp_label_paragraph(tbl_element):
    """The lettered client paragraph immediately above a PP block (or None)."""
    el = tbl_element.getprevious()
    hops = 0
    while el is not None and hops < 3:
        if el.tag.endswith("}p"):
            from docx.text.paragraph import Paragraph
            p = Paragraph(el, None)
            if "".join(r.text for r in p.runs).strip():
                return el
        elif el.tag.endswith("}tbl"):
            return None
        el = el.getprevious()
        hops += 1
    return None


def _sync_past_performance(doc, past_performance: list, report: Report) -> None:
    if not past_performance:
        return
    blocks = _pp_blocks(doc)
    if not blocks:
        for pp in past_performance:
            report.flag("Past Performance",
                        "No existing block to clone: add this engagement manually.",
                        KIND_ADD, new=str(pp.get("client", "")))
        return

    def find_block(nclient, nproject):
        exact = [t for t, c, p in blocks if c == nclient and p == nproject]
        if exact:
            return exact[0]
        same_client = [t for t, c, _p in blocks if c == nclient]
        if not nproject:
            # A project-less store record can't disambiguate between a client's
            # blocks -- match the first rather than ever appending a duplicate.
            return same_client[0] if same_client else None
        # Named project with no exact match: a genuinely new engagement
        # (append), even when the client already has other blocks.
        return None

    import copy as _copy
    for pp in past_performance:
        client = str(pp.get("client", "")).strip()
        if not client:
            continue
        project = str(pp.get("project", "") or "")
        tbl = find_block(_norm_name(client), _norm_name(project))
        label = f"Past Performance '{client}'" + (f" / '{project}'" if project else "")
        if tbl is not None:
            _fill_pp_block(tbl, pp, report, label, update_only=True)
            continue
        # Append a new block: clone the last block's table (+ its lettered
        # client paragraph, whose list numbering continues automatically).
        model_tbl, _c, _p = blocks[-1]
        new_tbl_el = _copy.deepcopy(model_tbl._tbl)
        model_tbl._tbl.addnext(new_tbl_el)
        label_el = _pp_label_paragraph(model_tbl._tbl)
        if label_el is not None:
            from docx.text.paragraph import Paragraph
            new_label_el = _copy.deepcopy(label_el)
            new_tbl_el.addprevious(new_label_el)
            para = Paragraph(new_label_el, model_tbl._parent)
            if para.runs:
                para.runs[0].text = client
                for r in para.runs[1:]:
                    r.text = ""
            else:
                para.add_run(client)
        new_tbl = Table(new_tbl_el, model_tbl._parent)
        _fill_pp_block(new_tbl, pp, report, label, update_only=False)
        blocks.append((new_tbl, _norm_name(client), _norm_name(project)))
        report.applied("Past Performance", "added engagement block from store", new=client)
